"""
导出软著申请用源代码文档。

中国软件著作权登记通常要求提交：
  - 源程序前 30 页 + 后 30 页（每页约 50 行，共约 3000 行）
  - 每页页眉标注软件名称与版本号

本脚本生成：
  1. full_source_code.txt       — 完整源程序清单（带文件分隔符）
  2. soft_copyright_source.txt  — 前 30 页 + 后 30 页（可直接复制到 Word）
  3. soft_copyright_source.docx   — Word 格式（若已安装 python-docx）
  4. file_manifest.txt          — 导出文件清单
"""

from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUT = ROOT / "exports" / "soft_copyright"

SOFTWARE_NAME = "纤维混凝土抗裂性能预测系统"
SOFTWARE_VERSION = "V1.0"
LINES_PER_PAGE = 50
FRONT_PAGES = 30
BACK_PAGES = 30

EXCLUDE_DIRS = {
    "outputs",
    "models",
    "data",
    ".git",
    "__pycache__",
    ".streamlit",
    "node_modules",
    ".venv",
    "venv",
    "exports",
}
EXCLUDE_FILES = {"_ui_theme_tmp.py", "test.py", "export_soft_copyright.py"}

# 核心模块优先，便于前 30 页展示主业务逻辑
PRIORITY_ORDER = [
    "app.py",
    "run_web.py",
    "api/main.py",
    "evaluate.py",
    "train_with_group_split.py",
    "train_lab_strength_residual.py",
    "train_thermal_stress_residual.py",
    "shap_analysis.py",
    "run_literature_pipeline.py",
]


def should_include(path: Path) -> bool:
    try:
        rel = path.relative_to(ROOT)
    except ValueError:
        return False
    for part in rel.parts:
        if part in EXCLUDE_DIRS:
            return False
    if path.name in EXCLUDE_FILES:
        return False
    if path.name.startswith("_") and "scripts" in rel.parts:
        return False
    return path.suffix == ".py"


def sort_key(path: Path) -> tuple:
    rel = path.relative_to(ROOT).as_posix()
    for i, name in enumerate(PRIORITY_ORDER):
        if rel == name or rel.endswith("/" + name):
            return (0, i, rel)
    if rel.startswith("src/"):
        return (1, rel)
    if rel.startswith("api/"):
        return (2, rel)
    if rel.startswith("experiments/"):
        return (3, rel)
    if rel.startswith("scripts/"):
        return (4, rel)
    return (5, rel)


def collect_files() -> list[Path]:
    files = [p for p in ROOT.rglob("*.py") if should_include(p)]
    return sorted(files, key=sort_key)


def build_line_stream(files: list[Path]) -> list[tuple[str, str]]:
    """返回 (source_file, line_text) 列表。"""
    stream: list[tuple[str, str]] = []
    for fp in files:
        text = fp.read_text(encoding="utf-8", errors="replace")
        rel = fp.relative_to(ROOT).as_posix()
        stream.append((rel, f"# ===== FILE: {rel} ====="))
        for line in text.splitlines():
            stream.append((rel, line.rstrip("\r\n")))
        stream.append((rel, ""))
    return stream


def paginate_lines(
    lines: list[str],
    *,
    software_name: str,
    version: str,
    lines_per_page: int,
) -> list[str]:
    """将代码行分页，每页带页眉。"""
    pages: list[str] = []
    total_pages = max(1, (len(lines) + lines_per_page - 1) // lines_per_page)
    for page_idx in range(total_pages):
        start = page_idx * lines_per_page
        end = start + lines_per_page
        chunk = lines[start:end]
        header = f"{software_name}  {version}    第 {page_idx + 1} 页 / 共 {total_pages} 页"
        pages.append(header)
        pages.extend(chunk)
        # 不足行数时用空行补齐至 lines_per_page（软著排版惯例）
        pad = lines_per_page - len(chunk)
        if pad > 0:
            pages.extend([""] * pad)
        pages.append("")  # 页间空行
    return pages


def extract_front_back(
    all_code_lines: list[str],
    *,
    lines_per_page: int,
    front_pages: int,
    back_pages: int,
) -> list[str]:
    front_count = front_pages * lines_per_page
    back_count = back_pages * lines_per_page
    total = len(all_code_lines)
    if total <= front_count + back_count:
        return all_code_lines
    return all_code_lines[:front_count] + all_code_lines[-back_count:]


def write_txt(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def write_docx(
    path: Path,
    pages_text: str,
    *,
    software_name: str,
    version: str,
) -> bool:
    try:
        from docx import Document
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import Cm, Pt
    except ImportError:
        return False

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(9)

    for block in pages_text.split("\n\n"):
        block = block.strip("\n")
        if not block:
            continue
        lines = block.splitlines()
        if not lines:
            continue
        header = lines[0]
        body = lines[1:]

        p_hdr = doc.add_paragraph(header)
        p_hdr.paragraph_format.space_after = Pt(6)
        run = p_hdr.runs[0]
        run.bold = True
        run.font.size = Pt(9)

        for line in body:
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.name = "Courier New"
                r.font.size = Pt(9)

    doc.core_properties.title = f"{software_name} 源程序"
    doc.core_properties.subject = f"{software_name} {version} 软著源代码"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return True


def main() -> None:
    parser = argparse.ArgumentParser(description="导出软著申请用源代码")
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT, help="输出目录")
    parser.add_argument("--name", default=SOFTWARE_NAME, help="软件全称")
    parser.add_argument("--version", default=SOFTWARE_VERSION, help="版本号")
    parser.add_argument("--lines-per-page", type=int, default=LINES_PER_PAGE)
    parser.add_argument("--front-pages", type=int, default=FRONT_PAGES)
    parser.add_argument("--back-pages", type=int, default=BACK_PAGES)
    args = parser.parse_args()

    files = collect_files()
    stream = build_line_stream(files)
    code_lines = [line for _, line in stream]

    out_dir = args.out
    out_dir.mkdir(parents=True, exist_ok=True)

    # 完整源程序
    full_parts = []
    for fp in files:
        rel = fp.relative_to(ROOT).as_posix()
        full_parts.append(f"{'=' * 72}")
        full_parts.append(f"  文件: {rel}")
        full_parts.append(f"{'=' * 72}")
        full_parts.append(fp.read_text(encoding="utf-8", errors="replace").rstrip())
        full_parts.append("")
    write_txt(out_dir / "full_source_code.txt", "\n".join(full_parts))

    # 前 30 + 后 30 页
    subset = extract_front_back(
        code_lines,
        lines_per_page=args.lines_per_page,
        front_pages=args.front_pages,
        back_pages=args.back_pages,
    )
    paginated = paginate_lines(
        subset,
        software_name=args.name,
        version=args.version,
        lines_per_page=args.lines_per_page,
    )
    submission_text = "\n".join(paginated)
    write_txt(out_dir / "soft_copyright_source.txt", submission_text)

    docx_ok = write_docx(
        out_dir / "soft_copyright_source.docx",
        submission_text,
        software_name=args.name,
        version=args.version,
    )

    manifest_lines = [
        f"软件名称: {args.name}",
        f"版本号: {args.version}",
        f"导出日期: {date.today().isoformat()}",
        f"源文件数量: {len(files)}",
        f"源代码总行数: {len(code_lines)}",
        f"提交用代码行数: {len(subset)} (前{args.front_pages}页+后{args.back_pages}页, 每页{args.lines_per_page}行)",
        "",
        "导出文件:",
        "  - full_source_code.txt      完整源程序（按文件分隔）",
        "  - soft_copyright_source.txt 软著提交用：前30页+后30页",
        "  - soft_copyright_source.docx" + ("  (已生成)" if docx_ok else "  (未生成，需 pip install python-docx)"),
        "  - file_manifest.txt         本清单",
        "",
        "源文件列表:",
    ]
    for fp in files:
        n = len(fp.read_text(encoding="utf-8", errors="replace").splitlines())
        manifest_lines.append(f"  {n:5d}  {fp.relative_to(ROOT).as_posix()}")

    write_txt(out_dir / "file_manifest.txt", "\n".join(manifest_lines))

    print(f"导出完成: {out_dir}")
    print(f"  源文件: {len(files)} 个, 共 {len(code_lines)} 行")
    print(f"  提交用: soft_copyright_source.txt ({args.front_pages}+{args.back_pages} 页)")
    if docx_ok:
        print(f"  Word:   soft_copyright_source.docx")
    else:
        print("  Word:   跳过 (请安装 python-docx 后重试)")


if __name__ == "__main__":
    main()
