import hashlib
import json
import math
import uuid
from datetime import datetime
from pathlib import Path

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st

from src.trust_engine import assess_trust
from src.features import FEATURE_COLUMNS, STRENGTH_GRADE_TO_MPA
from src.lab_experiment import (
    LOADING_COMPRESSION,
    LOADING_FLEXURAL,
    SPECIMEN_TYPES,
    estimate_strengths,
)
from src.mechanism import (
    feature_label,
    load_background_scaled,
    load_feature_importance_json,
    load_cv_metrics_json,
    load_training_metrics_json,
    local_ice_curve,
    local_ice_curve_proba,
    partial_dependence_1d,
    shap_risk_bar,
    synthetic_background_around,
    top_feature_importance_rows,
    top_feature_names,
)
from experiments.fiber_engineering.derive import derive_fiber_engineering_features
from experiments.environment_engineering.derive import derive_environment_engineering_features
from experiments.thermal_stress.derive import (
    derive_thermal_stress_features,
    thermal_stress_explain_sentence_zh,
)
from experiments.thermal_stress.engineering_chain import (
    derive_thermal_engineering_display,
    thermal_engineering_conclusion_zh,
)
from src.input_defaults import normalize_prediction_inputs
from src.model_performance_panel import render_model_performance_evaluation
from src.thermal_stress_inputs import series_for_thermal_derive
from src.paths import MODELS_DIR, OUTPUTS_DIR, PROJECT_ROOT

# 主开裂离线指标展示口径（Tab④ 等；仅文案，不参与推理）
CRACK_OFFLINE_METRICS_DISCLAIMER = (
    "以下为**历史 hold-out 指标**，**不代表当前输入实时误差**。"
    "当前页面加载权重与 `training_metrics.json` 为**弱同源**；"
    "JSON 中**缺少完整训练溯源字段**（如 `data_path`、`trained_at`、`model_version`）。"
)
from src.ui_theme import (
    COLORS,
    PLOTLY_CONFIG,
    apply_chart_theme,
    crack_formation_path_html,
    engineering_conclusion_html,
    engineering_conclusion_compact_html,
    engineering_explain_card_html,
    environment_driver_analysis_html,
    environment_engineering_kpi_html,
    environment_engineering_summary_panel_html,
    engineering_info_panel_html,
    fiber_bridge_cards_html,
    fiber_engineering_kpi_html,
    fiber_engineering_summary_panel_html,
    fiber_thermal_coordination_note_html,
    home_kpi_dashboard_html,
    kpi_card_html,
    mechanical_context_line_html,
    mechanical_summary_html,
    mechanism_conclusion_html,
    mechanism_flow_html,
    risk_tone_from_level,
    section_title_html,
    shap_narrative_html,
    style_gauge_indicator,
    tab_heading_html,
    thermal_engineering_chain_html,
    thermal_engineering_conclusion_banner_html,
    thermal_formula_step_html,
    thermal_full_engineering_flow_html,
    thermal_main_model_relation_html,
    thermal_status_chips_html,
    thermal_variable_cards_html,
    trust_notice_card_html,
    trust_conclusion_banner_html,
    trust_pipeline_methodology_html,
    trust_score_chip_html,
    input_range_check_panel_html,
    stability_status_panel_html,
    standards_compact_html,
    research_status_overview_html,
    research_module_header_html,
    governance_tier_strip_html,
)

# 机理页展示用简称（不改 FEATURE_COLUMNS / mechanism 算法）
_FEATURE_LABEL_PLAIN: dict[str, str] = {
    "w_b_ratio": "水胶比",
    "fiber_content": "纤维掺量",
    "aspect_ratio": "长径比",
    "mixing_water": "用水量",
    "binder_content": "胶材用量",
    "cement_content": "水泥用量",
    "cube_strength_mpa": "立方体强度",
    "tensile_strength": "纤维抗拉强度",
    "curing_days": "养护龄期",
    "temperature": "环境温度",
    "humidity": "环境湿度",
    "fiber_content_x_aspect_ratio": "掺量×长径比",
}

_ENGINEERING_FEATURE_EXPLAIN: dict[str, tuple[str, list[str], str]] = {
    "w_b_ratio": (
        "为什么水胶比影响开裂？",
        [
            "水胶比提高后，浆体孔隙率往往增加，干燥收缩与自收缩增大；",
            "约束条件下变形不协调时，更容易在薄弱区形成微裂缝并扩展；",
            "在模型中通常表现为裂缝宽度与开裂风险概率同步上升。",
        ],
        "工程上宜结合胶材组成与养护制度综合控制水胶比。",
    ),
    "fiber_content": (
        "为什么纤维掺量影响开裂？",
        [
            "纤维掺量影响裂纹桥接与残余抗拉能力，掺量不足时约束裂缝开展的能力偏弱；",
            "掺量过高若分散不良，也可能削弱基体均匀性；",
            "模型中纤维相关特征常对风险概率与裂缝密度有显著贡献。",
        ],
        "建议以试验验证分散性与最优掺量区间。",
    ),
    "mixing_water": (
        "为什么用水量影响开裂？",
        [
            "用水量偏高会增加浆体体积与收缩变形潜力；",
            "在胶材用量不变时，有效水胶比上升，早期收缩与开裂敏感性提高；",
            "模型中用水量常与裂缝宽度、风险等级同向变化。",
        ],
        "宜与减水剂、胶材用量一并优化。",
    ),
    "aspect_ratio": (
        "为什么长径比影响开裂？",
        [
            "长径比改变纤维架桥几何与应力传递路径；",
            "过长或过短都可能影响分散与握裹，进而改变抗裂贡献；",
            "模型通过掺量×长径比等组合特征捕捉非线性效应。",
        ],
        "应结合纤维外形与施工工艺解读。",
    ),
}

# session_state：预测结果历史（输入参数与预测概览）
SFC_PRED_HISTORY_KEY = "sfc_pred_history"
SFC_PRED_HIST_SIG_KEY = "sfc_detail_hist_sig"
SFC_PRED_HISTORY_MAX = 50


@st.cache_data(show_spinner=False)
def _cached_cube_strength_mpa_min_max_from_data_csv() -> tuple[float, float] | None:
    """
    从仓库 data 下训练用 CSV 读取 cube_strength_mpa 列的观测 min/max，
    用于「可能超出训练数据常见强度区」提示；读取失败则返回 None（不编造阈值）。
    """
    for name in ("training_data.csv", "training_data.example.csv"):
        p = PROJECT_ROOT / "data" / name
        if not p.is_file():
            continue
        try:
            df = pd.read_csv(p, nrows=200_000)
        except (OSError, UnicodeDecodeError, ValueError, pd.errors.ParserError):
            continue
        if "cube_strength_mpa" not in df.columns:
            continue
        s = pd.to_numeric(df["cube_strength_mpa"], errors="coerce").dropna()
        if s.empty:
            continue
        return float(s.min()), float(s.max())
    return None


@st.cache_data(show_spinner=False)
def _cached_main_model_input_ranges() -> dict[str, tuple[float, float]]:
    """主开裂训练 CSV 观测 min/max；缺列时回退 REFERENCE_RANGES。"""
    from src.lab_strength_residual.training_data_gate import REFERENCE_RANGES

    out: dict[str, tuple[float, float]] = dict(REFERENCE_RANGES)
    cols = (
        "w_b_ratio",
        "fiber_content",
        "aspect_ratio",
        "temperature",
        "humidity",
        "mixing_water",
    )
    for name in ("training_data.csv", "training_data.example.csv"):
        p = PROJECT_ROOT / "data" / name
        if not p.is_file():
            continue
        try:
            df = pd.read_csv(p, nrows=200_000)
        except (OSError, UnicodeDecodeError, ValueError, pd.errors.ParserError):
            continue
        for col in cols:
            if col not in df.columns:
                continue
            s = pd.to_numeric(df[col], errors="coerce").dropna()
            if len(s) >= 3:
                out[col] = (float(s.min()), float(s.max()))
        break
    return out


class Visualizer:
    def _plotly(self, fig: go.Figure, key: str | None = None) -> None:
        """绘制 Plotly 图；key 供「高级分析」多图同页时避免 Streamlit 组件 id 冲突。"""
        kwargs: dict = {"use_container_width": True, "config": PLOTLY_CONFIG}
        if key is not None:
            kwargs["key"] = key
        st.plotly_chart(fig, **kwargs)
    @staticmethod
    def _is_bad_num(v) -> bool:
        if v is None:
            return True
        if isinstance(v, float):
            return not math.isfinite(v)
        return False

    @staticmethod
    def _is_good_num(v) -> bool:
        """有限数值（int/float 可转且非 NaN/Inf）；None 或非数值型为 False。"""
        if v is None:
            return False
        try:
            x = float(v)
        except (TypeError, ValueError):
            return False
        return math.isfinite(x)

    @staticmethod
    def _fmt_value(v):
        if Visualizer._is_bad_num(v):
            return "—"
        if isinstance(v, float):
            return f"{v:.4f}"
        if isinstance(v, (dict, list)):
            return str(v)
        return str(v)

    @staticmethod
    def _word_safe_text(s: object, max_len: int = 30000) -> str:
        """去除 Word/OpenXML 不允许的控制字符，避免写入 docx 失败。"""
        if s is None:
            return ""
        t = str(s)
        t = "".join(c for c in t if ord(c) >= 32 or c in "\t\n\r")
        return t[:max_len]

    def _table_grid_safe(self, table) -> None:
        for name in ("Table Grid", "Light Grid Accent 1", "Medium Grid 1 Accent 1", "Table Normal"):
            try:
                table.style = name
                return
            except (ValueError, KeyError, AttributeError):
                continue

    @staticmethod
    def _fmt_float(v, decimals: int, suffix: str = "") -> str:
        if Visualizer._is_bad_num(v):
            return "—"
        try:
            return f"{float(v):.{decimals}f}{suffix}"
        except (TypeError, ValueError):
            return "—"

    @staticmethod
    def _prediction_snapshot_sig(user_inputs: dict, flat_display: dict) -> str:
        """用于判断是否为同一次预测快照（参数+结果概览）。"""
        payload = json.dumps(
            {"u": user_inputs, "p": flat_display},
            sort_keys=True,
            ensure_ascii=False,
            default=str,
        )
        return hashlib.sha256(payload.encode("utf-8")).hexdigest()

    @staticmethod
    def _deep_copy_jsonable(d: dict) -> dict:
        return json.loads(json.dumps(d, default=str))

    def _sync_prediction_history(self, user_inputs: dict, flat_display: dict) -> None:
        """当本次预测与上次不同时，在 history 头部插入一条快照。"""
        sig = self._prediction_snapshot_sig(user_inputs, flat_display)
        hist: list = st.session_state.setdefault(SFC_PRED_HISTORY_KEY, [])
        prev = st.session_state.get(SFC_PRED_HIST_SIG_KEY)
        if prev == sig:
            return
        st.session_state[SFC_PRED_HIST_SIG_KEY] = sig
        rec = {
            "id": uuid.uuid4().hex[:12],
            "ts": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "user_inputs": self._deep_copy_jsonable(dict(user_inputs)),
            "flat_display": self._deep_copy_jsonable(dict(flat_display)),
        }
        hist.insert(0, rec)
        while len(hist) > SFC_PRED_HISTORY_MAX:
            hist.pop()

    @staticmethod
    def _try_load_json_dict(path: Path) -> dict | None:
        """
        只读容错加载 JSON 对象。
        中文说明：用于展示离线训练报告；文件不存在、非 dict、或解析失败时返回 None，不向界面编造指标。
        """
        try:
            if not path.is_file():
                return None
            with open(path, encoding="utf-8") as f:
                raw = json.load(f)
            return raw if isinstance(raw, dict) else None
        except (OSError, json.JSONDecodeError, TypeError, UnicodeDecodeError):
            return None

    @staticmethod
    def _live_prediction_reliability_lines(result: dict) -> list[str]:
        """
        归纳「本次输入的实时预测」路径说明（数据全部来自 predict_all 已有键；不修改 predictor）。
        中文说明：与下方离线 JSON 指标严格区分。
        """
        lines: list[str] = []
        mid = result.get("intermediate") or {}
        preds = result.get("predictions", {}) or {}
        sd = preds.get("state_dimension") or {}
        td = preds.get("time_dimension") or {}

        lines.append(
            "本标签页中的 **P、w、裂缝密度** 等为 **已加载主模型的本次输入实时预测**；"
            "**不是** `outputs/training_metrics.json` 中的离线测试集数值。"
        )
        dens_src = mid.get("crack_density_source")
        if dens_src == "fallback":
            lines.append(
                "裂缝密度：本次推理走 **经验回退**（`intermediate.crack_density_source=fallback`），"
                "即未采用密度回归器直接数值。"
            )
        elif dens_src == "regressor":
            lines.append(
                "裂缝密度：本次由 **密度回归模型** 输出（`crack_density_source=regressor`）。"
            )
        else:
            cn = sd.get("crack_density_source_cn")
            if isinstance(cn, str) and cn.strip():
                lines.append(f"裂缝密度来源说明：{cn}")
            else:
                lines.append(
                    "裂缝密度：返回结构中未给出细分来源字段；请以页内数值及机理分析为准。"
                )

        if isinstance(td.get("note"), str) and td["note"].strip():
            lines.append(
                f"时间维度附注：{td['note']}（**启发式/信息性**，非标准试验测定。）"
            )

        if isinstance(sd.get("stress_strength_note_cn"), str) and sd[
            "stress_strength_note_cn"
        ].strip():
            lines.append("应力-强度比 σ/fₜ：为 **启发式量级估计**，非 GB 试验直接输出。")

        lines.append(
            "「**试验估算**」标签页中的抗压/抗折：为 **国标公式基线（信息性）**，"
            "与主开裂三维度 **不同管线**，请勿混读为同一模型输出。"
        )
        return lines

    @staticmethod
    def _lab_strength_offline_markdown(rep: dict | None) -> str:
        """
        从 lab_strength_residual_report.json 提取可读摘要（仅展示文件中已有字段）。
        """
        if not rep:
            return (
                "未找到 `outputs/lab_strength/lab_strength_residual_report.json`。"
                "**请先训练并生成报告**（例如项目根目录执行 "
                "`py train_lab_strength_residual.py ...`）。**此处不展示任何编造指标。**"
            )

        chunks: list[str] = []
        data_p = rep.get("data")
        if data_p is not None:
            chunks.append(f"- 报告内记录的数据源：`{data_p}`")

        n_final = rep.get("n_rows_final_used")
        if n_final is not None:
            chunks.append(f"- 入模有效行数（报告统计）：**{n_final}**")

        for flag_key, label in (
            ("practical_gain", "practical_gain（报告布尔字段）"),
            ("residual_not_recommended", "residual_not_recommended（报告布尔字段）"),
        ):
            if flag_key in rep:
                chunks.append(f"- {label}：`{rep.get(flag_key)}`")

        dm = rep.get("default_method_by_task") or {}
        for task_key, task_cn in (
            ("compressive", "抗压"),
            ("flexural", "抗折"),
        ):
            info = dm.get(task_key)
            if not isinstance(info, dict):
                chunks.append(f"- **{task_cn}** 默认策略：报告中无 `default_method_by_task.{task_key}`。")
                continue
            strat = info.get("strategy", "—")
            desc = info.get("description") or ""
            learner = info.get("residual_learner")
            learner_s = "无（仅公式）" if learner is None else str(learner)
            chunks.append(
                f"- **{task_cn}** 默认策略：`{strat}`；残差学习器：`{learner_s}`"
                + (f"；{desc}" if desc else "")
            )

            task_block = rep.get(task_key)
            if not isinstance(task_block, dict):
                continue
            cv_fm = task_block.get("cv_fold_metrics") or {}
            n_samp = cv_fm.get("n_samples")
            if n_samp is not None:
                chunks.append(
                    f"  - 离线 CV 样本量 **n_samples={n_samp}**（与本次侧栏输入无关）。"
                )

            oof = task_block.get("oof_global_metrics") or {}
            method_key = "formula_only"
            if isinstance(learner, str) and learner in oof:
                method_key = learner
            elif isinstance(learner, str) and learner:
                method_key = "formula_only"
                chunks.append(
                    f"  - 提示：默认残差学习器 `{learner}` 在 `oof_global_metrics` 中无键，"
                    f"OOF 数值改读 **`formula_only`**。"
                )

            metrics = oof.get(method_key) if isinstance(oof, dict) else {}
            if not isinstance(metrics, dict):
                metrics = {}
            mae = metrics.get("mae")
            rmse = metrics.get("rmse")
            r2 = metrics.get("r2")
            if (
                isinstance(mae, (int, float))
                and isinstance(rmse, (int, float))
                and isinstance(r2, (int, float))
                and math.isfinite(float(mae))
                and math.isfinite(float(rmse))
                and math.isfinite(float(r2))
            ):
                chunks.append(
                    f"  - 全局 **OOF**（方法 `{method_key}`）：MAE={float(mae):.6f}，"
                    f"RMSE={float(rmse):.6f}，R²={float(r2):.6f}（**离线交叉验证汇总**，非当前输入误差）。"
                )
            else:
                chunks.append(
                    f"  - 未在报告中找到 `{method_key}` 的完整 OOF mae/rmse/r2 数值，**不编造**。"
                )

        return "\n".join(chunks)

    @staticmethod
    def _crack_main_offline_markdown(tm: dict | None) -> str:
        """
        从 outputs/training_metrics.json 提取主开裂模型离线测试集摘要（字段以文件为准）。
        """
        if not tm:
            return (
                "未找到 `outputs/training_metrics.json`。"
                "**请先完成主模型训练**（例如 `py -m src.train_model`）后生成。"
                "**此处不展示任何编造指标。**"
            )

        chunks: list[str] = [CRACK_OFFLINE_METRICS_DISCLAIMER, ""]
        n_train = tm.get("n_train")
        n_test = tm.get("n_test")
        if n_train is not None and n_test is not None:
            chunks.append(
                f"- 划分：**n_train={n_train}**，**n_test={n_test}**"
                "（**历史 hold-out**，与当前单次预测无关）。"
            )

        for title, key, mode in (
            ("裂缝宽度 crack_width", "crack_width", "reg"),
            ("裂缝密度 crack_density", "crack_density", "reg"),
            ("开裂风险 cracking_risk", "cracking_risk", "cls"),
        ):
            block = tm.get(key)
            if not isinstance(block, dict):
                continue
            if mode == "reg":
                mae = block.get("test_mae")
                rmse = block.get("test_rmse")
                r2 = block.get("test_r2")
                if all(
                    isinstance(x, (int, float)) and math.isfinite(float(x))
                    for x in (mae, rmse, r2)
                ):
                    chunks.append(
                        f"- **{title}** 测试集：MAE={float(mae):.6f}，RMSE={float(rmse):.6f}，"
                        f"R²={float(r2):.6f}（**历史 hold-out**）。"
                    )
                else:
                    chunks.append(f"- **{title}**：缺少完整 test_mae/test_rmse/test_r2，不展示。")
            else:
                acc = block.get("test_accuracy")
                f1m = block.get("test_macro_f1")
                wf1 = block.get("test_weighted_f1")
                parts = []
                if isinstance(acc, (int, float)) and math.isfinite(float(acc)):
                    parts.append(f"accuracy={float(acc):.6f}")
                if isinstance(f1m, (int, float)) and math.isfinite(float(f1m)):
                    parts.append(f"macro_f1={float(f1m):.6f}")
                if isinstance(wf1, (int, float)) and math.isfinite(float(wf1)):
                    parts.append(f"weighted_f1={float(wf1):.6f}")
                if parts:
                    chunks.append(
                        f"- **{title}** 测试集：{', '.join(parts)}（**历史 hold-out**）。"
                    )
                else:
                    chunks.append(f"- **{title}**：无可解析分类指标字段。")

        return "\n".join(chunks) if chunks else "- 文件存在但无可解析指标字段。"

    def show_input_range_warnings(self, user_inputs: dict) -> None:
        """
        阶段 C：侧栏关键量超出闸门参考带或本地强度观测区间时给短提示；
        仅 UI，不拦截预测；数值阈值仍只复用 REFERENCE_RANGES 与 data CSV。
        """
        from src.lab_strength_residual.training_data_gate import REFERENCE_RANGES

        # 强度观测区间：读不到时弱提示，避免用户以为已做分布外检查
        bounds = _cached_cube_strength_mpa_min_max_from_data_csv()
        if bounds is None:
            st.caption("未读取到本地强度观测区间，当前未启用强度分布外提醒。")

        items: list[str] = []
        for key, label_short in (
            ("w_b_ratio", "水胶比"),
            ("fiber_content", "纤维掺量"),
            ("aspect_ratio", "长径比"),
        ):
            tup = REFERENCE_RANGES.get(key)
            if not isinstance(tup, (list, tuple)) or len(tup) < 2:
                continue
            lo, hi = float(tup[0]), float(tup[1])
            raw = user_inputs.get(key)
            try:
                v = float(raw)
            except (TypeError, ValueError):
                continue
            if not math.isfinite(v):
                continue
            if v < lo or v > hi:
                items.append(
                    f"· **{label_short}** {v:g} 超出常用带 [{lo:g}, {hi:g}]，"
                    f"可能偏训练分布外 — 预测仅供参考。"
                )

        grade = user_inputs.get("strength_grade")
        if (
            bounds is not None
            and isinstance(grade, str)
            and grade in STRENGTH_GRADE_TO_MPA
        ):
            mpa = float(STRENGTH_GRADE_TO_MPA[grade])
            lo_b, hi_b = bounds
            if mpa < lo_b - 1e-9 or mpa > hi_b + 1e-9:
                items.append(
                    f"· **强度** {grade}（≈{mpa:g} MPa）超出本地表观测 [{lo_b:g}, {hi_b:g}]，"
                    f"可能偏训练分布外 — 预测仅供参考。"
                )

        if not items:
            return
        st.warning("**输入检查** · 不拦截预测\n\n" + "\n".join(items))

    def _overview_lab_formula_snapshot(self, user_inputs: dict):
        """Tab1 用：默认 150mm 立方体 + 标准加载的信息性力学基线（非主开裂模型）。"""
        grade = str(user_inputs.get("strength_grade", "C30")).strip().upper()
        try:
            fcu = float(STRENGTH_GRADE_TO_MPA[grade])
        except (KeyError, TypeError, ValueError):
            fcu = 30.0
        try:
            fc_pct = float(user_inputs.get("fiber_content", 1.0))
        except (TypeError, ValueError):
            fc_pct = 1.0
        return estimate_strengths(
            "立方体（边长可变）",
            cube_edge_mm=150.0,
            prism_b_mm=150.0,
            prism_h_mm=150.0,
            prism_l_mm=300.0,
            beam_b_mm=100.0,
            beam_h_mm=100.0,
            beam_span_mm=400.0,
            loading_compression=LOADING_COMPRESSION[0],
            loading_flexural=LOADING_FLEXURAL[0],
            cube_strength_mpa=fcu,
            fiber_content_pct=fc_pct,
            compute_flexural=True,
        )

    @staticmethod
    def _overview_summary_sentence(result: dict) -> str:
        preds = result.get("predictions", {}) or {}
        sd = preds.get("state_dimension") or {}
        rd = preds.get("risk_dimension") or {}
        lvl = rd.get("alert_level", preds.get("risk_level", "—"))
        w = sd.get("crack_width_mm", preds.get("crack_width"))
        dens = sd.get("crack_density_per_m2", preds.get("crack_density"))
        p = sd.get("risk_probability", preds.get("risk_confidence"))
        parts = [f"开裂风险预测模型给出 **{lvl}** 预警"]
        if p is not None:
            try:
                parts.append(f"风险概率约 **{float(p) * 100:.0f}%**")
            except (TypeError, ValueError):
                pass
        if w is not None:
            parts.append(f"预测裂缝宽度约 **{float(w):.3f} mm**")
        if dens is not None:
            parts.append(f"裂缝密度约 **{float(dens):.2f} 条/m²**")
        parts.append(
            "力学 MPa 为公式基线示意；温度应力为解释指数，均不替代规范试验与主模型综合判定。"
        )
        return "；".join(parts) + "。"

    def _build_engineering_conclusion(
        self, result: dict, user_inputs: dict
    ) -> tuple[str, list[str], str]:
        """基于现有预测与输入的启发式工程摘要（仅展示，无新算法）。"""
        preds = result.get("predictions", {}) or {}
        sd = preds.get("state_dimension") or {}
        rd = preds.get("risk_dimension") or {}
        level = str(rd.get("alert_level", preds.get("risk_level", "未知")))
        factors: list[str] = []

        try:
            w_b = float(user_inputs.get("w_b_ratio", 0.4))
            if w_b > 0.42:
                factors.append("水胶比偏高")
        except (TypeError, ValueError):
            pass

        merged = normalize_prediction_inputs(user_inputs)
        feats = derive_thermal_stress_features(series_for_thermal_derive(merged))
        if self._thermal_tsi_computable(feats):
            tsi_f = float(feats["thermal_stress_index"])
            if self._thermal_tsi_risk_band(tsi_f) in ("中", "高"):
                factors.append("温度应力指数偏高")

        wref = sd.get("crack_width_gb50010") or {}
        cmp_cn = str(wref.get("comparison_cn", ""))
        if any(k in cmp_cn for k in ("超过", "接近", "偏大", "偏高")):
            factors.append("裂缝宽度接近或超过常用验算参照限值")

        try:
            fc = float(user_inputs.get("fiber_content", 1.0))
            if fc < 0.8:
                factors.append("钢纤维掺量偏低")
        except (TypeError, ValueError):
            pass

        try:
            p = float(sd.get("risk_probability", preds.get("risk_confidence", 0.5)))
            if p > 0.7 and "风险概率" not in "".join(factors):
                factors.append("模型风险概率处于高预警带")
        except (TypeError, ValueError):
            pass

        if "高" in level:
            lead = "当前配合比下模型预测开裂风险较高"
            advice = (
                "建议适当降低水胶比并优化纤维掺量，同时核查温差、约束与养护条件，"
                "必要时开展试配与现场监测。"
            )
        elif "中" in level:
            lead = "当前配合比下模型预测开裂风险处于中等水平"
            advice = (
                "建议关注水胶比、纤维掺量与温度—约束组合，结合规范验算与试件试验做复核。"
            )
        else:
            lead = "当前配合比下模型预测开裂风险相对较低"
            advice = (
                "可维持现行配合比思路，仍建议结合 GB 50010 裂缝验算与现场养护记录综合判定。"
            )

        if not factors:
            factors = ["综合材料与工艺特征共同作用于当前预警带"]

        return lead, factors, advice

    @staticmethod
    def _compact_engineering_reasons(factors: list[str]) -> list[str]:
        mapping = {
            "水胶比偏高": "水胶比偏高",
            "温度应力指数偏高": "温度应力指数偏高",
            "裂缝宽度接近或超过常用验算参照限值": "裂缝宽度接近经验限值",
            "钢纤维掺量偏低": "纤维掺量偏低",
            "模型风险概率处于高预警带": "风险概率进入高预警带",
            "综合材料与工艺特征共同作用于当前预警带": "材料与工艺综合影响当前预警带",
        }
        return [mapping.get(f, f) for f in factors[:2]]

    @staticmethod
    def _compact_engineering_suggestions(level: str, factors: list[str]) -> list[str]:
        items: list[str] = []
        if any("水胶" in f for f in factors):
            items.append("降低水胶比")
        if any("纤维" in f for f in factors):
            items.append("优化纤维掺量")
        if any("温度" in f for f in factors):
            items.append("核查温差与养护条件")
        if any("裂缝" in f for f in factors) and "降低水胶比" not in items:
            items.append("控制裂缝宽度发展")
        if not items:
            if "高" in level:
                items = ["降低水胶比", "优化纤维掺量", "核查温差与养护条件"]
            elif "中" in level:
                items = ["关注水胶比与纤维掺量", "结合规范验算复核"]
            else:
                items = ["维持现行配合比思路", "结合规范与养护记录判定"]
        return items[:3]

    def _crack_width_status_tag(self, result: dict) -> str:
        preds = result.get("predictions", {}) or {}
        sd = preds.get("state_dimension") or {}
        wref = sd.get("crack_width_gb50010") or {}
        cmp_cn = str(wref.get("comparison_cn", ""))
        if any(k in cmp_cn for k in ("超过", "接近")):
            return "接近经验限值"
        if any(k in cmp_cn for k in ("未超过", "低于", "安全", "满足")):
            return "处于安全区"
        try:
            w = float(sd.get("crack_width_mm", preds.get("crack_width", 0)))
            if w <= 0.25:
                return "处于安全区"
            if w >= 0.28:
                return "接近经验限值"
        except (TypeError, ValueError):
            pass
        return ""

    @staticmethod
    def _prob_warning_band(p: float) -> str:
        if p >= 0.7:
            return "高预警带"
        if p >= 0.3:
            return "中预警带"
        return "低预警带"

    @staticmethod
    def _risk_level_subtitle(level: str) -> str:
        if "高" in level:
            return "模型预测当前工况裂缝风险较高"
        if "中" in level:
            return "模型预测当前工况裂缝风险处于中等水平"
        return "模型预测当前工况裂缝风险相对较低"

    def show_crack_core_metrics(
        self, result: dict, user_inputs: dict | None = None
    ) -> None:
        """主开裂指标：1 主卡 + 2 辅卡 + 3 次级卡（HTML，不改预测值）。"""
        preds = result.get("predictions", {}) or {}
        sd = preds.get("state_dimension") or {}
        rd = preds.get("risk_dimension") or {}
        lvl = str(rd.get("alert_level", preds.get("risk_level", "未知")))
        tone = risk_tone_from_level(lvl)
        p_txt = self._fmt_float(
            sd.get("risk_probability", preds.get("risk_confidence")), 3
        )
        w_txt = self._fmt_float(
            sd.get("crack_width_mm", preds.get("crack_width")), 3
        )
        dens_txt = self._fmt_float(
            sd.get("crack_density_per_m2", preds.get("crack_density")), 2
        )

        est = None
        if user_inputs is not None:
            est = self._overview_lab_formula_snapshot(user_inputs)
        comp_txt = (
            f"{est.compressive_formula_pred:.2f}"
            if est is not None
            else "—"
        )
        flex_txt = (
            f"{est.flexural_formula_pred:.2f}"
            if est is not None and math.isfinite(est.flexural_formula_pred)
            else "—"
        )


        width_tag = self._crack_width_status_tag(result)
        if width_tag == "处于安全区":
            width_tag = "安全范围"

        try:
            p_raw = float(sd.get("risk_probability", preds.get("risk_confidence", 0.0)))
        except (TypeError, ValueError):
            p_raw = 0.0
        if not math.isfinite(p_raw):
            p_raw = 0.0
        prob_band = self._prob_warning_band(min(max(p_raw, 0.0), 1.0))

        st.markdown(
            home_kpi_dashboard_html(
                risk_level=lvl,
                risk_tone=tone,
                risk_subtitle=self._risk_level_subtitle(lvl),
                prob_value=p_txt,
                prob_band=prob_band,
                width_value=w_txt,
                width_tag=width_tag,
                density_value=dens_txt,
                density_unit="条/m²",
                compressive_value=comp_txt,
                flexural_value=flex_txt,
            ),
            unsafe_allow_html=True,
        )

    def show_offline_training_metrics_panel(self) -> None:
        """离线训练指标（历史数据，非当前样本实时误差）。"""
        lab_json_path = PROJECT_ROOT / "outputs" / "lab_strength" / "lab_strength_residual_report.json"
        crack_tm_path = PROJECT_ROOT / "outputs" / "training_metrics.json"
        st.caption(CRACK_OFFLINE_METRICS_DISCLAIMER)
        with st.expander("历史 hold-out 指标明细（JSON 摘要）", expanded=False):
            lab_rep = self._try_load_json_dict(lab_json_path)
            tm_rep = self._try_load_json_dict(crack_tm_path)
            st.markdown("##### 抗压/抗折强度残差管线（lab_strength_residual）")
            st.markdown(self._lab_strength_offline_markdown(lab_rep))
            st.markdown("---")
            st.markdown("##### 开裂风险预测模型（outputs/training_metrics.json）")
            st.markdown(self._crack_main_offline_markdown(tm_rep))
            st.caption(f"路径：`{lab_json_path}` · `{crack_tm_path}`")

    def show_standards_disclaimer_panel(self) -> None:
        with st.expander("标准依据与完整免责声明", expanded=False):
            st.markdown("- GB/T 50081-2019《普通混凝土力学性能试验方法标准》")
            st.markdown("- GB 50010-2010（2015 年版）《混凝土结构设计规范》")
            st.markdown("- GB/T 50082-2009《普通混凝土长期性能和耐久性能试验方法标准》")
            st.markdown("- 《大体积混凝土施工技术规程》")
            st.markdown("- GB 50164-2011《混凝土质量控制标准》")
            st.markdown("- JGJ/T 221-2010《纤维混凝土应用技术规程》")
            st.caption(
                "本系统不能替代结构设计、施工验收、试验检测及监理依据中的完整合规判定；"
                "不宣传“高精度”，超出训练分布或协议不一致时结果仅供参考。"
            )

    def show_data_governance_status_panel(self) -> None:
        st.markdown("##### 数据层级与治理状态（简述）")
        st.caption(
            "**A 类样本**用于主 OOF；**B 类**用于技术预览；**C 类**与 synthetic **不进入主结论**。"
        )
        gov_path = PROJECT_ROOT / "outputs" / "crack_governance" / "crack_training_governance.json"
        phase_path = PROJECT_ROOT / "outputs" / "crack_governance" / "PHASE_STATUS.md"
        rep = self._try_load_json_dict(gov_path)
        if rep:
            tc = rep.get("tier_ABC_hold_counts") or {}
            st.markdown(
                f"- 主开裂训练表行数：**{rep.get('row_count', '—')}**  \n"
                f"- 暂缓：**{tc.get('hold_pending', '—')}** · "
                f"A 类候选：**{tc.get('tier_A_candidate', 0)}** · "
                f"B：**{tc.get('tier_B', 0)}** · "
                f"C/非法 tier：**{tc.get('tier_C_or_illegal_tier', 0)}**"
            )
            ga = rep.get("group_audit") or {}
            st.markdown(
                f"- 已填 `source_group` 行数：**{ga.get('n_rows_with_nonempty_source_group', 0)}**"
            )
        else:
            st.caption(f"未找到 `{gov_path}`，请先运行治理诊断脚本。")
        if phase_path.is_file():
            st.caption(f"阶段说明：`{phase_path}`")

    def show_tab_comprehensive_prediction(self, result: dict, user_inputs: dict) -> None:
        """Tab1：预测结果（普通用户主路径）。"""
        st.markdown(tab_heading_html("预测结果"), unsafe_allow_html=True)
        _lead, factors, _advice = self._build_engineering_conclusion(result, user_inputs)
        preds = result.get("predictions", {}) or {}
        rd = preds.get("risk_dimension") or {}
        judgment = str(rd.get("alert_level", preds.get("risk_level", "未知")))
        st.markdown(
            engineering_conclusion_compact_html(
                judgment,
                self._compact_engineering_reasons(factors),
                self._compact_engineering_suggestions(judgment, factors),
            ),
            unsafe_allow_html=True,
        )
        st.caption(
            "以下为**开裂风险预测模型**实时输出；视觉层级：风险等级 > P > 裂缝宽度 > 辅助指标。"
        )
        self.show_crack_core_metrics(result, user_inputs)
        self.show_risk_gauge(result)
        st.caption(self._overview_summary_sentence(result))

    @staticmethod
    def _short_loading_label(label: str) -> str:
        s = str(label or "—")
        if "（" in s:
            return s.split("（", 1)[0].strip()
        return s

    def _render_lab_session_controls(self, user_inputs: dict) -> dict:
        """试件/加载控件（仅 UI 状态，计算仍走 estimate_strengths）。"""
        st.radio(
            "估算范围",
            ("仅抗压", "抗压与抗折"),
            index=1,
            key="lab_estimate_scope",
            horizontal=True,
        )
        compute_flex = (
            st.session_state.get("lab_estimate_scope", "抗压与抗折") == "抗压与抗折"
        )
        spec = st.selectbox("试件类型", SPECIMEN_TYPES, key="lab_specimen")
        load_c = st.selectbox(
            "抗压试验 · 加载方式", LOADING_COMPRESSION, key="lab_load_c"
        )
        load_f: str | None = None
        if compute_flex:
            load_f = st.selectbox(
                "抗折试验 · 加载方式", LOADING_FLEXURAL, key="lab_load_f"
            )

        cube_edge = 150.0
        pb, ph, pl = 150.0, 150.0, 300.0
        bb, bh, bspan = 100.0, 100.0, 400.0

        if spec == "立方体（边长可变）":
            cube_edge = float(
                st.number_input(
                    "立方体边长 a (mm)",
                    min_value=50.0,
                    max_value=300.0,
                    value=150.0,
                    step=5.0,
                    key="lab_cube_a",
                )
            )
        elif spec == "棱柱体（轴心抗压）":
            a1, a2, a3 = st.columns(3)
            with a1:
                pb = float(
                    st.number_input(
                        "宽度 b (mm)",
                        min_value=50.0,
                        max_value=300.0,
                        value=150.0,
                        step=5.0,
                        key="lab_pb",
                    )
                )
            with a2:
                ph = float(
                    st.number_input(
                        "高度 h (mm)",
                        min_value=50.0,
                        max_value=500.0,
                        value=150.0,
                        step=5.0,
                        key="lab_ph",
                    )
                )
            with a3:
                pl = float(
                    st.number_input(
                        "长度 L (mm)",
                        min_value=100.0,
                        max_value=600.0,
                        value=300.0,
                        step=5.0,
                        key="lab_pl",
                    )
                )
        else:
            a1, a2, a3 = st.columns(3)
            with a1:
                bb = float(
                    st.number_input(
                        "梁宽 b (mm)",
                        min_value=40.0,
                        max_value=200.0,
                        value=100.0,
                        step=5.0,
                        key="lab_bb",
                    )
                )
            with a2:
                bh = float(
                    st.number_input(
                        "梁高 h (mm)",
                        min_value=80.0,
                        max_value=300.0,
                        value=100.0,
                        step=5.0,
                        key="lab_bh",
                    )
                )
            with a3:
                bspan = float(
                    st.number_input(
                        "跨度 L (mm)",
                        min_value=200.0,
                        max_value=1200.0,
                        value=400.0,
                        step=10.0,
                        key="lab_span",
                    )
                )

        grade = str(user_inputs.get("strength_grade", "C30")).strip().upper()
        try:
            fc_pct = float(user_inputs.get("fiber_content", 1.0))
        except (TypeError, ValueError):
            fc_pct = 1.0

        return {
            "spec": spec,
            "load_c": load_c,
            "load_f": load_f,
            "compute_flex": compute_flex,
            "cube_edge": cube_edge,
            "pb": pb,
            "ph": ph,
            "pl": pl,
            "bb": bb,
            "bh": bh,
            "bspan": bspan,
            "grade": grade,
            "fc_pct": fc_pct,
        }

    def _lab_estimate_bundle(
        self, user_inputs: dict, ctrl: dict
    ) -> tuple[object, dict, dict]:
        from src.lab_strength_residual.lab_mix_features import (
            lab_mix_extra_dict_from_user_inputs,
        )

        est = estimate_strengths(
            ctrl["spec"],
            cube_edge_mm=ctrl["cube_edge"],
            prism_b_mm=ctrl["pb"],
            prism_h_mm=ctrl["ph"],
            prism_l_mm=ctrl["pl"],
            beam_b_mm=ctrl["bb"],
            beam_h_mm=ctrl["bh"],
            beam_span_mm=ctrl["bspan"],
            loading_compression=ctrl["load_c"],
            loading_flexural=ctrl["load_f"],
            cube_strength_mpa=float(
                STRENGTH_GRADE_TO_MPA.get(ctrl["grade"], 30.0)
            ),
            fiber_content_pct=ctrl["fc_pct"],
            compute_flexural=ctrl["compute_flex"],
        )
        wr_extra = lab_mix_extra_dict_from_user_inputs(user_inputs)
        est.detail = dict(est.detail)
        est.detail["lab_strength_lab_mix_extra"] = wr_extra
        meta = self._lab_display_meta(ctrl, user_inputs)
        return est, wr_extra, meta

    def _lab_display_meta(self, ctrl: dict, user_inputs: dict) -> dict:
        spec = ctrl["spec"]
        if spec == "立方体（边长可变）":
            geom = f"{ctrl['cube_edge']:.0f} mm 立方体"
            spec_short = f"{ctrl['cube_edge']:.0f} mm 立方体"
        elif spec == "棱柱体（轴心抗压）":
            geom = (
                f"b×h×L = {ctrl['pb']:.0f}×{ctrl['ph']:.0f}×{ctrl['pl']:.0f} mm"
            )
            spec_short = "棱柱体（轴心抗压）"
        else:
            geom = (
                f"梁 b×h×L = {ctrl['bb']:.0f}×{ctrl['bh']:.0f}×{ctrl['bspan']:.0f} mm"
            )
            spec_short = "梁式试件（抗折）"

        scope = "抗压与抗折" if ctrl["compute_flex"] else "仅抗压"
        load_f_disp = (
            self._short_loading_label(ctrl["load_f"])
            if ctrl["compute_flex"] and ctrl["load_f"]
            else "—"
        )
        try:
            w_b = float(user_inputs.get("w_b_ratio", 0.4))
        except (TypeError, ValueError):
            w_b = 0.4

        return {
            "spec": spec,
            "spec_short": spec_short,
            "geometry": geom,
            "load_c": self._short_loading_label(ctrl["load_c"]),
            "load_f": load_f_disp,
            "scope": scope,
            "compute_flex": ctrl["compute_flex"],
            "grade": ctrl["grade"],
            "fc_pct": ctrl["fc_pct"],
            "w_b_ratio": w_b,
        }

    @staticmethod
    def _lab_specimen_panel_html(meta: dict) -> str:
        rows = [
            ("试件类型", meta["spec_short"]),
            ("几何尺寸", meta["geometry"]),
            ("抗压加载", meta["load_c"]),
            ("抗折加载", meta["load_f"]),
            ("估算范围", meta["scope"]),
        ]
        return engineering_info_panel_html(
            "当前试件参数",
            rows,
            note="以上为当前试验设置摘要；如需修改，请展开页顶「调整试验设置」。",
        )

    def _water_reducer_panel_html(self, user_inputs: dict, wr_extra: dict) -> str:
        wr_type = str(user_inputs.get("water_reducer_type", "无"))
        wr_unknown = bool(user_inputs.get("water_reduction_rate_unknown", False))
        wr_rate = float(wr_extra.get("water_reduction_rate_pct", -1.0))
        wr_adj_miss = float(wr_extra.get("adjusted_w_b_ratio_missing_flag", 1.0))
        wr_adj = float(wr_extra.get("adjusted_w_b_ratio", -1.0))
        try:
            w_b = float(user_inputs.get("w_b_ratio", 0.4))
        except (TypeError, ValueError):
            w_b = 0.4

        if wr_unknown:
            rate_txt = "未知"
            adj_txt = "暂不折算（减水率未知）"
        elif wr_rate < 0:
            rate_txt = "—"
            adj_txt = f"{w_b:.3f}（未折算）"
        else:
            rate_txt = f"{wr_rate:.0f}%"
            adj_txt = (
                f"{wr_adj:.3f}"
                if wr_adj_miss < 0.5
                else f"{w_b:.3f}（未折算）"
            )

        rows = [
            ("减水剂", wr_type),
            ("减水率", rate_txt),
            ("修正后水胶比", adj_txt),
        ]
        return engineering_info_panel_html(
            "减水剂修正状态",
            rows,
            note=(
                "当前修正链已进入后端 lab_strength 派生逻辑，"
                "但当前页 MPa 结果仍以 GB 公式基线为主，不随修正水胶比自动改写。"
            ),
        )

    def _mechanical_trust_bullets(self, user_inputs: dict) -> list[str]:
        bullets = [
            "当前结果主要基于 GB/T 50081 与 GB 50010 公式基线，并结合历史训练数据作辅助说明",
            "结果适用于工程辅助评估与方案比选",
            "不替代正式试验检测与标准检测报告",
        ]
        caution = self._mechanical_input_range_caution(user_inputs)
        if caution:
            bullets.insert(1, caution)
        else:
            bullets.insert(
                1,
                "当前侧栏关键参数位于常见工程参考范围内",
            )
        return bullets

    def _mechanical_input_range_caution(self, user_inputs: dict) -> str | None:
        from src.lab_strength_residual.training_data_gate import REFERENCE_RANGES

        issues: list[str] = []
        for key, label_short in (
            ("w_b_ratio", "水胶比"),
            ("fiber_content", "纤维掺量"),
            ("aspect_ratio", "长径比"),
        ):
            tup = REFERENCE_RANGES.get(key)
            if not isinstance(tup, (list, tuple)) or len(tup) < 2:
                continue
            lo, hi = float(tup[0]), float(tup[1])
            try:
                v = float(user_inputs.get(key))
            except (TypeError, ValueError):
                continue
            if math.isfinite(v) and (v < lo or v > hi):
                issues.append(label_short)

        if not issues:
            return None
        return (
            f"侧栏「{'、'.join(issues)}」可能偏离常见训练参考带，当前力学估算仅供参考"
        )

    def _show_mechanical_results_header(
        self, est, meta: dict, user_inputs: dict
    ) -> None:
        comp = f"{est.compressive_formula_pred:.2f}"
        if meta["compute_flex"] and math.isfinite(est.flexural_formula_pred):
            flex = f"{est.flexural_formula_pred:.2f}"
        else:
            flex = "—"

        bullets = [f"抗压强度估算约 {comp} MPa"]
        if flex != "—":
            bullets.append(f"抗折强度估算约 {flex} MPa")
        else:
            bullets.append("当前设置为仅抗压，未计算抗折强度")
        bullets.extend(
            [
                "当前结果基于 GB 公式基线与历史训练数据辅助说明",
                "不替代标准试验报告",
            ]
        )
        st.markdown(
            mechanical_summary_html(
                bullets,
                footer="以上为信息性估算，正式判定请以标准试件试验为准。",
            ),
            unsafe_allow_html=True,
        )

        flex_val = flex
        kpi_row = (
            kpi_card_html(
                label="抗压强度",
                value=comp,
                unit="MPa",
                tier="primary",
                tone="comp",
                icon="🧱",
                hint="GB/T 50081 公式基线",
            )
            + kpi_card_html(
                label="抗折强度",
                value=flex_val,
                unit="MPa" if flex_val != "—" else "",
                tier="primary",
                tone="flex",
                icon="🧱",
                hint="GB 50010 公式基线" if flex_val != "—" else "未启用抗折估算",
            )
        )
        st.markdown(
            f'<div class="sfc-mech-kpi-row">{kpi_row}</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            mechanical_context_line_html(
                meta["grade"],
                f"{meta['fc_pct']:.1f}%",
            ),
            unsafe_allow_html=True,
        )

    def _show_mechanical_advanced_analysis(
        self, user_inputs: dict, est, wr_extra: dict
    ) -> None:
        """研发向指标与 JSON，默认折叠。"""
        with st.expander("高级分析（开发/研究）", expanded=False):
            st.caption(
                "以下为历史训练、管线诊断与中间量，**不影响**当前试件公式估算结果。"
            )
            with st.expander("字段作用与公式路径说明", expanded=False):
                st.markdown(
                    """
**抗压公式（当前估算）** 主要随：试件类型、几何尺寸、抗压加载方式、强度等级映射的 fcu,k。

**抗折公式** 在启用「抗压与抗折」时参与：试件类型、梁尺寸与跨度、抗折加载方式、fcu,k、纤维掺量。

侧栏多数配合比字段用于**开裂风险主模型**；减水剂相关项在后端生成派生量，**本页 MPa 仍以 GB 公式基线为主**，请勿将修正水胶比与页顶 MPa 直接等同。

抗折加载方式**不改变**抗压结果（抗压路径不读取该字段）。
                    """.strip()
                )

            lab_json_path = (
                PROJECT_ROOT / "outputs" / "lab_strength" / "lab_strength_residual_report.json"
            )
            lab_rep = self._try_load_json_dict(lab_json_path)

            with st.expander("历史训练交叉验证与默认策略", expanded=False):
                self._render_lab_strength_offline_metrics_body()

            with st.expander("离线报告文字摘要", expanded=False):
                st.markdown(self._lab_strength_offline_markdown(lab_rep))

            with st.expander("计算中间量（JSON）", expanded=False):
                st.json(est.detail)
                for note in est.notes:
                    st.caption(note)

            with st.expander("减水剂派生量明细（技术字段）", expanded=False):
                st.json(wr_extra)

            with st.expander("异常行与人工复核记录", expanded=False):
                drop_path = (
                    PROJECT_ROOT
                    / "outputs"
                    / "lab_strength"
                    / "lab_strength_dropped_rows.json"
                )
                rep_path = lab_json_path
                dropped = self._try_load_json_dict(drop_path)
                if dropped:
                    st.json(dropped)
                else:
                    st.caption(f"未找到 {drop_path}")
                rep = self._try_load_json_dict(rep_path)
                if rep:
                    st.caption(
                        f"manual_review 行数（报告）："
                        f"{rep.get('manual_review_rows_count', '—')}"
                    )

    def show_tab_mechanical_support(self, user_inputs: dict) -> None:
        """Tab2：力学性能（工程评估主路径，研发指标默认折叠）。"""
        st.markdown(tab_heading_html("力学性能"), unsafe_allow_html=True)
        st.markdown(section_title_html("力学性能评估"), unsafe_allow_html=True)
        with st.expander("调整试验设置", expanded=False):
            st.caption("修改试件尺寸、加载方式或估算范围；结果区将自动更新。")
            ctrl = self._render_lab_session_controls(user_inputs)
        est, wr_extra, meta = self._lab_estimate_bundle(user_inputs, ctrl)
        self._show_mechanical_results_header(est, meta, user_inputs)
        st.markdown(section_title_html("试件与加载参数"), unsafe_allow_html=True)
        st.markdown(
            self._lab_specimen_panel_html(meta),
            unsafe_allow_html=True,
        )
        st.markdown(section_title_html("减水剂修正状态"), unsafe_allow_html=True)
        st.markdown(
            self._water_reducer_panel_html(user_inputs, wr_extra),
            unsafe_allow_html=True,
        )
        st.markdown(
            trust_notice_card_html(
                "结果可信度说明",
                self._mechanical_trust_bullets(user_inputs),
            ),
            unsafe_allow_html=True,
        )
        self._show_mechanical_advanced_analysis(user_inputs, est, wr_extra)

    @staticmethod
    def _feat_label_plain(name: str) -> str:
        if name in _FEATURE_LABEL_PLAIN:
            return _FEATURE_LABEL_PLAIN[name]
        lab = feature_label(name)
        if "(" in lab:
            return lab.split("(", 1)[0].strip()
        return lab

    @staticmethod
    def _mechanism_risk_level(result: dict) -> str:
        preds = result.get("predictions", {}) or {}
        rd = preds.get("risk_dimension") or {}
        return str(rd.get("alert_level", preds.get("risk_level", "未知")))

    def _prepare_mechanism_context(
        self, predictor, X_df: pd.DataFrame, pred_result: dict
    ) -> dict:
        """准备 SHAP / 背景样本等（仅读取已有模型，不改算法）。"""
        mid = pred_result.get("intermediate") or {}
        pred_class = int(mid.get("risk_class_argmax", 0))
        imp_all = load_feature_importance_json(MODELS_DIR) or {}
        Xo = X_df[FEATURE_COLUMNS].astype(np.float64, copy=False)
        x_scaled = predictor.scaler.transform(Xo)
        x_single = x_scaled[0:1]
        X_bg = load_background_scaled(predictor.scaler)
        bg_synthetic = False
        if X_bg is None:
            X_bg = synthetic_background_around(x_single, n=220)
            bg_synthetic = True
        shp_full = None
        try:
            if predictor.crack_risk_model is not None:
                shp_full = shap_risk_bar(
                    predictor.crack_risk_model,
                    X_bg,
                    x_single,
                    pred_class,
                    top_k=14,
                )
        except Exception:
            shp_full = None
        return {
            "pred_class": pred_class,
            "imp_all": imp_all,
            "X_bg": X_bg,
            "x_single": x_single,
            "bg_synthetic": bg_synthetic,
            "shp": shp_full,
        }

    def _thermal_feats_for_mechanism(self, user_inputs: dict) -> dict:
        merged = normalize_prediction_inputs(user_inputs)
        return derive_thermal_stress_features(series_for_thermal_derive(merged))

    def _build_mechanism_driver_bullets(
        self,
        user_inputs: dict,
        result: dict,
        ctx: dict,
        feats: dict,
    ) -> list[str]:
        drivers: list[str] = []
        seen: set[str] = set()

        def add(text: str) -> None:
            if text and text not in seen and len(drivers) < 4:
                seen.add(text)
                drivers.append(text)

        shp = ctx.get("shp")
        if shp:
            names, values = shp
            for n, v in sorted(
                zip(names, values), key=lambda x: -float(x[1])
            )[:4]:
                if float(v) > 0:
                    add(self._driver_phrase_for_feature(n, user_inputs, "raise"))

        try:
            w_b = float(user_inputs.get("w_b_ratio", 0.4))
            if w_b > 0.42:
                add("水胶比较高")
        except (TypeError, ValueError):
            pass

        try:
            fc = float(user_inputs.get("fiber_content", 1.0))
            if fc < 0.8:
                add("纤维约束能力不足")
            elif fc > 2.2:
                add("纤维掺量偏高需关注分散")
        except (TypeError, ValueError):
            pass

        try:
            water = float(user_inputs.get("mixing_water", 180))
            if water > 190:
                add("用水量偏高")
        except (TypeError, ValueError):
            pass

        if not self._thermal_tsi_computable(feats):
            add("温度应力路径未闭合")
        elif self._thermal_tsi_risk_band(float(feats["thermal_stress_index"])) in (
            "中",
            "高",
        ):
            add("温度应力指数偏高")

        if not drivers:
            imp = (ctx.get("imp_all") or {}).get("cracking_risk") or {}
            if imp:
                top = top_feature_names(imp, n=3)
                for n in top:
                    add(f"{self._feat_label_plain(n)}为关键影响因素")
            else:
                add("材料与工艺特征共同作用于当前预警带")

        return drivers[:4]

    def _driver_phrase_for_feature(
        self, feat: str, user_inputs: dict, direction: str
    ) -> str:
        plain = self._feat_label_plain(feat)
        if feat == "w_b_ratio" and direction == "raise":
            return "水胶比较高"
        if feat == "fiber_content":
            try:
                fc = float(user_inputs.get("fiber_content", 1))
                if direction == "raise" and fc > 1.5:
                    return "纤维掺量偏高"
                if direction == "raise" and fc < 0.8:
                    return "纤维约束能力不足"
            except (TypeError, ValueError):
                pass
            return f"{plain}推高风险" if direction == "raise" else f"{plain}降低风险"
        if feat == "mixing_water" and direction == "raise":
            return "用水量偏高"
        if feat == "aspect_ratio" and direction == "raise":
            return "长径比处于不利区间"
        if direction == "raise":
            return f"{plain}偏高"
        return f"{plain}有利于降低风险"

    def _mechanism_synthesis_sentence(self, level: str) -> str:
        if "高" in level:
            return (
                "材料收缩变形与约束能力不匹配，"
                "导致裂缝宽度与裂缝密度上升。"
            )
        if "中" in level:
            return (
                "材料变形与边界约束存在一定不协调，"
                "裂缝开展风险处于中等水平。"
            )
        return (
            "当前配合比下材料—约束匹配相对较好，"
            "裂缝开展倾向较低。"
        )

    def _build_crack_path_nodes(
        self, user_inputs: dict, result: dict, feats: dict
    ) -> list[dict]:
        level = self._mechanism_risk_level(result)
        risk_tone = risk_tone_from_level(level)

        try:
            w_b = float(user_inputs.get("w_b_ratio", 0.4))
            wb_status = "偏高 ↑" if w_b > 0.42 else ("偏低 ↓" if w_b < 0.32 else "适中")
            wb_tone = "high" if w_b > 0.42 else ("low" if w_b < 0.35 else "info")
        except (TypeError, ValueError):
            wb_status, wb_tone = "—", "muted"

        try:
            fc = float(user_inputs.get("fiber_content", 1.0))
            if fc < 0.8:
                fiber_status, fiber_tone = "偏弱", "high"
            elif fc > 2.0:
                fiber_status, fiber_tone = "偏高", "mid"
            else:
                fiber_status, fiber_tone = "一般", "info"
        except (TypeError, ValueError):
            fiber_status, fiber_tone = "—", "muted"

        if not self._thermal_tsi_computable(feats):
            thermal_status, thermal_tone = "路径缺失", "muted"
        else:
            band = self._thermal_tsi_risk_band(float(feats["thermal_stress_index"]))
            thermal_status = f"{band}风险"
            thermal_tone = (
                "high" if band == "高" else ("mid" if band == "中" else "low")
            )

        preds = result.get("predictions", {}) or {}
        sd = preds.get("state_dimension") or {}
        w_mm = sd.get("crack_width_mm", preds.get("crack_width"))
        try:
            w_f = float(w_mm)
            crack_status = (
                f"约 {w_f:.3f} mm"
                if math.isfinite(w_f)
                else "—"
            )
        except (TypeError, ValueError):
            crack_status = "—"

        return [
            {
                "icon": "🧱",
                "title": "材料与环境输入",
                "status": f"水胶比 {wb_status}",
                "tone": wb_tone,
            },
            {
                "icon": "🌡️",
                "title": "收缩/温差变形",
                "status": thermal_status,
                "tone": thermal_tone,
            },
            {
                "icon": "🔩",
                "title": "约束与刚度",
                "status": f"纤维 {fiber_status}",
                "tone": fiber_tone,
            },
            {
                "icon": "⚡",
                "title": "局部拉应力",
                "status": "累积偏高" if risk_tone == "high" else "可控",
                "tone": risk_tone if risk_tone != "neutral" else "info",
            },
            {
                "icon": "📉",
                "title": "裂缝扩展",
                "status": crack_status,
                "tone": "mid" if risk_tone == "mid" else risk_tone,
            },
            {
                "icon": "⚠️",
                "title": "开裂风险",
                "status": level,
                "tone": risk_tone if risk_tone != "neutral" else "mid",
            },
        ]

    def _shap_contribution_narrative(
        self, names: list[str], values: np.ndarray
    ) -> str:
        inc = [
            (n, v) for n, v in zip(names, values) if float(v) > 0
        ]
        dec = [
            (n, v) for n, v in zip(names, values) if float(v) < 0
        ]
        inc.sort(key=lambda x: -x[1])
        dec.sort(key=lambda x: x[1])
        parts: list[str] = []
        if dec:
            labs = "、".join(self._feat_label_plain(n) for n, _ in dec[:2])
            parts.append(f"{labs}降低了风险")
        if inc:
            labs = "、".join(self._feat_label_plain(n) for n, _ in inc[:2])
            parts.append(f"较高的{labs}提高了裂缝概率")
        if not parts:
            return "当前模型认为：各材料与工艺因素在当前预警带内共同作用。"
        return "当前模型认为：" + "，而".join(parts) + "。"

    def _engineering_explain_cards_html(
        self, ctx: dict, user_inputs: dict
    ) -> str:
        pick: list[str] = []
        shp = ctx.get("shp")
        if shp:
            names, values = shp
            order = np.argsort(-np.abs(values))[:6]
            pick = [names[i] for i in order if names[i] in _ENGINEERING_FEATURE_EXPLAIN]
        if len(pick) < 2:
            imp = (ctx.get("imp_all") or {}).get("cracking_risk") or {}
            for n in top_feature_names(imp, n=8):
                if n in _ENGINEERING_FEATURE_EXPLAIN and n not in pick:
                    pick.append(n)
                if len(pick) >= 2:
                    break
        cards = []
        for feat in pick[:2]:
            title, bullets, foot = _ENGINEERING_FEATURE_EXPLAIN[feat]
            cards.append(engineering_explain_card_html(title, bullets, foot))
        return "".join(cards) if cards else ""

    _THERMAL_ENGINEERING_VAR_CARDS: list[dict[str, str]] = [
        {
            "title": "ΔT（温差）",
            "meaning": "构件内外或控制工况下的有效温差，驱动自由热变形需求。",
            "unit": "℃（与 K 等价差值）",
            "direction": "↑ → 温度应变与温度应力通常增大",
        },
        {
            "title": "α（热膨胀系数）",
            "meaning": "单位温升下的线变形能力，反映材料热胀冷缩幅度。",
            "unit": "1/℃（界面常按 ×10⁻⁶/℃ 输入）",
            "direction": "↑ → 同样温差下温度变形需求增大",
        },
        {
            "title": "E（弹性模量）",
            "meaning": "材料刚度；约束下协调变形越困难，拉应力解释项越高。",
            "unit": "MPa（可由侧栏输入或由强度等级经验估算）",
            "direction": "↑ → 同样变形不协调时应力水平倾向增大",
        },
        {
            "title": "R（约束系数）",
            "meaning": "边界、厚度、岩体/钢筋约束与滑移层等综合约束程度。",
            "unit": "无量纲（0~1 量级解释系数）",
            "direction": "↑ → 自由变形越难释放，温度应力增大",
        },
        {
            "title": "ε_T（温度应变）",
            "meaning": "若无约束时的自由热应变量级，σ=E·ε 的变形基础。",
            "unit": "无量纲应变",
            "direction": "↑ → 后续温度应力解释项增大",
        },
        {
            "title": "σ_T*（解释应力标量）",
            "meaning": "按 R·E·ε 量级的工程示意应力，用于理解约束效应。",
            "unit": "MPa 量级（标注为解释代理量，非 FE 求解）",
            "direction": "↑ → 拉应力不利程度增加",
        },
        {
            "title": "f_t（抗拉代理）",
            "meaning": "由立方体抗压强度按 GB 50010 表 4.1.3-2 映射的轴心抗拉标准值。",
            "unit": "MPa（工程解释代理量）",
            "direction": "↑ → 同等 σ_T* 下 η 降低，抗裂裕度增大",
        },
        {
            "title": "η（应力-抗拉能力比）",
            "meaning": "温度拉应力解释量与抗拉代理之比，开裂倾向判据。",
            "unit": "无量纲",
            "direction": "↑ → 温度因素主导开裂的风险升高",
        },
    ]

    @staticmethod
    def _fmt_thermal_sci(x: float | None, *, scale: float = 1.0, unit: str = "") -> str:
        if x is None or not math.isfinite(float(x)):
            return "—"
        v = float(x) * scale
        if abs(v) >= 1000 or (abs(v) < 0.01 and v != 0):
            s = f"{v:.3e}"
        else:
            s = f"{v:.4g}"
        return f"{s}{unit}"

    def _thermal_formula_steps_html(self, display: dict, merged: dict) -> str:
        feats = display.get("feats") or {}
        dt_ok = display.get("step1_ok")
        s2_ok = display.get("step2_ok")
        s3_ok = display.get("step3_ok")

        alpha = display.get("alpha_per_C")
        dt = display.get("delta_T_eff")
        eps = display.get("epsilon_T")
        E = display.get("E_MPa")
        R = display.get("restraint_factor_R")
        sig = display.get("sigma_T_explain")
        ft = display.get("f_t_proxy_MPa")
        eta = display.get("eta")
        ft_src_zh = display.get("f_t_proxy_source_zh") or "工程解释代理量"

        e_note = (
            "侧栏弹性模量"
            if merged.get("elastic_modulus_E_user") is not None
            else "由强度等级估算 E"
        )

        step1 = thermal_formula_step_html(
            1,
            "温度应变",
            ["ε_T = α × ΔT"],
            "温度变化会引起材料自由热膨胀或收缩；此处取有效温差绝对值参与应变量级。",
            [
                ("α", self._fmt_thermal_sci(alpha, scale=1e6, unit=" ×10⁻⁶/℃") if alpha else "—"),
                ("ΔT", self._fmt_thermal_sci(dt, unit=" ℃")),
                ("ε_T", self._fmt_thermal_sci(eps) if eps is not None else "—"),
            ],
        )
        step2 = thermal_formula_step_html(
            2,
            "约束温度应力",
            ["σ_T* = R × E × α × ΔT", "σ_T* = R × E × ε_T"],
            "当结构受到约束时，自由温度变形无法完全释放，会形成温度拉应力解释量。",
            [
                ("R", self._fmt_thermal_sci(R) if s2_ok else "—"),
                ("E", f"{self._fmt_thermal_sci(E, unit=' MPa')}（{e_note}）" if E else "—"),
                ("α", self._fmt_thermal_sci(alpha, scale=1e6, unit=" ×10⁻⁶/℃") if alpha else "—"),
                ("ΔT", self._fmt_thermal_sci(dt, unit=" ℃")),
                ("σ_T*", self._fmt_thermal_sci(sig, unit=" MPa") if sig is not None else "—"),
            ],
            note="σ_T* 为工程解释应力标量，不强称为真实 MPa 或有限元结果。",
        )
        eta_band = display.get("eta_risk_band") or "—"
        step3 = thermal_formula_step_html(
            3,
            "开裂风险判据",
            ["η = σ_T* / f_t"],
            "将温度应力解释量与混凝土抗拉能力对比；"
            "f_t 优先实测劈裂抗拉，其次抗折，再次由强度等级经验估算。",
            [
                ("σ_T*", self._fmt_thermal_sci(sig, unit=" MPa") if sig is not None else "—"),
                (
                    "f_t",
                    f"{self._fmt_thermal_sci(ft, unit=' MPa')}（{ft_src_zh}）"
                    if ft is not None
                    else "—",
                ),
                ("η", self._fmt_thermal_sci(eta) if eta is not None else "—"),
                ("判据带", f"{eta_band}风险" if s3_ok else "—"),
            ],
            note=(
                "建议：η < 0.6 低风险；0.6~1.0 中风险；> 1.0 高风险。"
                "（仅用于温度因素解释，不替代主模型开裂风险概率。）"
            ),
        )
        if not dt_ok:
            step1 += '<p class="sfc-thermal-formula-desc">⚠ Step 1 未闭合：补全温差路径。</p>'
        if not s2_ok:
            step2 += '<p class="sfc-thermal-formula-desc">⚠ Step 2 未闭合：补全 α、E 或约束等级。</p>'
        if not s3_ok:
            step3 += (
                '<p class="sfc-thermal-formula-desc">⚠ Step 3 未闭合：'
                f"{str(display.get('eta_block_reason', '需强度等级以映射 f_t'))}。</p>"
            )
        return step1 + step2 + step3

    def show_thermal_stress_engineering_compact(
        self, user_inputs: dict, pred_result: dict | None = None
    ) -> None:
        """机理 Tab：工程公式驱动的温度应力解释链（Phase 1，不改主模型）。"""
        st.markdown(
            section_title_html("工程温度应力计算链（Phase 1）"),
            unsafe_allow_html=True,
        )
        merged = normalize_prediction_inputs(user_inputs)
        row = series_for_thermal_derive(merged)
        display = derive_thermal_engineering_display(row)
        feats = display.get("feats") or derive_thermal_stress_features(row)

        dt_ok = feats.get("delta_T_eff_missing_flag") == 0
        r_ok = feats.get("restraint_factor_R_missing_flag") == 0
        s1 = display.get("step1_ok")
        s2 = display.get("step2_ok")
        s3 = display.get("step3_ok")

        # —— 1. 工程结论 ——
        lead = thermal_engineering_conclusion_zh(display)
        meta_parts = []
        if s3 and display.get("eta") is not None:
            meta_parts.append(f"η ≈ {float(display['eta']):.3f}")
        if self._thermal_tsi_computable(feats):
            tsi_f = float(feats["thermal_stress_index"])
            meta_parts.append(
                f"温度应力指数 {tsi_f:.2f}（{self._thermal_tsi_risk_band(tsi_f)}，无量纲辅助量）"
            )
        meta = " · ".join(meta_parts) if meta_parts else "补全侧栏可选输入后可闭合完整公式链"
        st.markdown(
            thermal_engineering_conclusion_banner_html(
                lead,
                meta,
                band=str(display.get("eta_risk_band")) if s3 else None,
            ),
            unsafe_allow_html=True,
        )
        self._show_fiber_thermal_coordination_note(user_inputs)
        env_note = self._env_derive_row(user_inputs).get("environment_thermal_linkage_note_zh")
        if env_note:
            st.markdown(
                f'<div class="sfc-fiber-thermal-note">{env_note}</div>',
                unsafe_allow_html=True,
            )

        # —— 2. 工程公式链 ——
        st.markdown(section_title_html("工程公式链"), unsafe_allow_html=True)
        st.markdown(self._thermal_formula_steps_html(display, merged), unsafe_allow_html=True)

        opt_ctx = display.get("optional_context") or {}
        if self._thermal_optional_context_has_data(opt_ctx):
            st.markdown(section_title_html("可选输入（解释/验证）"), unsafe_allow_html=True)
            self._show_thermal_optional_context_panel(opt_ctx)

        # —— 3. 工程变量解释 ——
        st.markdown(section_title_html("工程变量说明"), unsafe_allow_html=True)
        st.markdown(
            thermal_variable_cards_html(self._THERMAL_ENGINEERING_VAR_CARDS),
            unsafe_allow_html=True,
        )

        # —— 4. 温度 → 应力 → 开裂流程图 ——
        st.markdown(section_title_html("温度 → 应力 → 开裂流程"), unsafe_allow_html=True)
        st.markdown(
            thermal_full_engineering_flow_html(
                dt_ok=bool(dt_ok),
                strain_ok=bool(s1),
                stress_ok=bool(s2),
                criterion_ok=bool(s3),
            ),
            unsafe_allow_html=True,
        )

        # —— 5. 与主模型关系 ——
        st.markdown(section_title_html("与主开裂模型输出的关系"), unsafe_allow_html=True)
        preds = (pred_result or {}).get("predictions") or {}
        cw = preds.get("crack_width")
        rp = preds.get("risk_probability")
        alert = preds.get("alert_level", "—")
        cw_s = f"{float(cw):.3f} mm" if self._is_good_num(cw) else "—"
        rp_s = f"P ≈ {float(rp):.1%}" if self._is_good_num(rp) else "—"
        tsi_line = (
            f"无量纲指数 ≈ {float(feats['thermal_stress_index']):.2f}"
            f"（{self._thermal_tsi_risk_band(float(feats['thermal_stress_index']))}）"
            if self._thermal_tsi_computable(feats)
            else "指数未闭合，仅展示已算得的公式步骤"
        )
        st.markdown(
            thermal_main_model_relation_html(cw_s, rp_s, str(alert), tsi_line),
            unsafe_allow_html=True,
        )

        # —— 6. 公式说明（底部） ——
        st.markdown(
            '<div class="sfc-thermal-footnote">'
            "本模块为<strong>工程解释型</strong>温度应力指数与公式链，"
            "<strong>不代表</strong>有限元热力耦合计算结果，"
            "主要用于解释温度变化与开裂倾向之间的关系。"
            "</div>",
            unsafe_allow_html=True,
        )

        # —— 7. 高级诊断（默认折叠） ——
        with st.expander("高级诊断（开发/研究）", expanded=False):
            self._show_thermal_advanced_diagnostics(merged, display, feats)

    def _show_risk_contribution_section(self, ctx: dict) -> None:
        st.markdown(section_title_html("当前样本风险贡献"), unsafe_allow_html=True)
        shp = ctx.get("shp")
        if not shp:
            st.caption(
                "当前无法生成风险贡献图；请查看下方「高级分析」中的特征重要性，"
                "或确认已安装 shap 且开裂风险模型已加载。"
            )
            imp = (ctx.get("imp_all") or {}).get("cracking_risk") or {}
            if imp:
                names = top_feature_names(imp, n=8)
                st.markdown(
                    "离线重要性较高的因素："
                    + "、".join(self._feat_label_plain(n) for n in names)
                )
            return

        names, values = shp
        show_n = min(8, len(names))
        names = names[:show_n]
        values = values[:show_n]
        labels = [self._feat_label_plain(n) for n in names]
        colors = np.where(values >= 0, COLORS["danger"], COLORS["primary_light"])
        st.markdown(
            shap_narrative_html(self._shap_contribution_narrative(names, values)),
            unsafe_allow_html=True,
        )
        fig_s = go.Figure(
            go.Bar(
                x=values,
                y=labels,
                orientation="h",
                marker=dict(color=colors, line=dict(width=0)),
            )
        )
        h_shap = max(320, 28 * len(labels))
        fig_s.update_layout(
            title="当前样本 · 风险贡献（红=提高，蓝=降低）",
            xaxis_title="相对贡献",
            height=h_shap,
            margin=dict(l=160, r=28, t=48, b=40),
        )
        apply_chart_theme(fig_s, height=h_shap)
        self._plotly(fig_s)
        if len(shp[0]) > show_n:
            with st.expander(f"查看更多因素（共 {len(shp[0])} 项）", expanded=False):
                rest_n = len(shp[0]) - show_n
                st.caption(
                    f"另有 {rest_n} 项因素贡献较小，完整列表见「高级分析」。"
                )

    def show_mechanism_engineering_surface(
        self,
        predictor,
        X_df: pd.DataFrame,
        pred_result: dict,
        user_inputs: dict,
    ) -> None:
        """机理 Tab 主路径：结论 → 路径 → SHAP → 工程卡片 → 温度。"""
        ctx = self._prepare_mechanism_context(predictor, X_df, pred_result)
        feats = self._thermal_feats_for_mechanism(user_inputs)
        level = self._mechanism_risk_level(pred_result)
        drivers = self._build_mechanism_driver_bullets(
            user_inputs, pred_result, ctx, feats
        )
        lead = (
            "当前高风险主要由："
            if "高" in level
            else "当前开裂倾向主要来自："
        )
        st.markdown(
            mechanism_conclusion_html(
                drivers,
                self._mechanism_synthesis_sentence(level),
                lead=lead,
            ),
            unsafe_allow_html=True,
        )
        st.markdown(section_title_html("开裂形成路径"), unsafe_allow_html=True)
        st.markdown(
            crack_formation_path_html(
                self._build_crack_path_nodes(user_inputs, pred_result, feats)
            ),
            unsafe_allow_html=True,
        )
        self._show_risk_contribution_section(ctx)
        self._show_fiber_engineering_bridge_section(user_inputs)
        self._show_environment_engineering_section(user_inputs)
        cards_html = self._engineering_explain_cards_html(ctx, user_inputs)
        if cards_html:
            st.markdown(section_title_html("工程意义解释"), unsafe_allow_html=True)
            st.markdown(cards_html, unsafe_allow_html=True)
        self.show_thermal_stress_engineering_compact(user_inputs, pred_result)

    def _fiber_derive_row(self, user_inputs: dict) -> dict:
        merged = normalize_prediction_inputs(user_inputs)
        return derive_fiber_engineering_features(pd.Series(merged))

    def _env_derive_row(self, user_inputs: dict) -> dict:
        merged = normalize_prediction_inputs(user_inputs)
        return derive_environment_engineering_features(pd.Series(merged))

    def _environment_kpi_tuples(self, feats: dict) -> list[tuple[str, str, str]]:
        summary = feats.get("environment_engineering_summary") or {}
        evap = summary.get("evaporation_risk_zh", "—").replace("蒸发风险：", "")
        tgrad = summary.get("thermal_gradient_risk_zh", "—").replace("温差风险：", "")
        curing = summary.get("curing_adequacy_zh", "—").replace("养护充分性：", "")
        surface = summary.get("surface_moisture_loss_zh", "—").replace("表层失水风险：", "")
        return [
            ("蒸发风险", evap, f"index={feats.get('evaporation_risk_index', '—')}"),
            ("温差风险", tgrad, f"index={feats.get('thermal_gradient_risk', '—')}"),
            ("养护水平", curing.split("（")[0] if "（" in curing else curing, "工程解释"),
            ("表层失水", surface, f"shrink={feats.get('surface_shrinkage_risk', '—')}"),
        ]

    def _show_environment_engineering_section(self, user_inputs: dict) -> None:
        """开裂机理 Tab：环境驱动开裂分析 + 工程状态卡。"""
        feats = self._env_derive_row(user_inputs)
        bullets = feats.get("environment_driver_bullets_zh") or []
        if bullets:
            st.markdown(section_title_html("环境驱动开裂分析"), unsafe_allow_html=True)
            st.markdown(environment_driver_analysis_html(bullets), unsafe_allow_html=True)
        st.markdown(section_title_html("环境场工程状态"), unsafe_allow_html=True)
        st.markdown(
            environment_engineering_kpi_html(self._environment_kpi_tuples(feats)),
            unsafe_allow_html=True,
        )

    def _show_environment_engineering_summary_advanced(self, user_inputs: dict) -> None:
        with st.expander("环境工程体系 · environment_engineering_summary", expanded=False):
            feats = self._env_derive_row(user_inputs)
            summary = feats.get("environment_engineering_summary") or {}
            st.markdown(
                environment_engineering_summary_panel_html(summary),
                unsafe_allow_html=True,
            )
            st.markdown("**派生量（解释层）**")
            rows = [
                ("evaporation_risk_index", feats.get("evaporation_risk_index")),
                ("thermal_gradient_risk", feats.get("thermal_gradient_risk")),
                ("surface_shrinkage_risk", feats.get("surface_shrinkage_risk")),
                ("evaporation_risk_score", feats.get("evaporation_risk_score")),
                ("thermal_gradient_risk_score", feats.get("thermal_gradient_risk_score")),
                ("curing_adequacy_score", feats.get("curing_adequacy_score")),
            ]
            st.dataframe(
                pd.DataFrame(rows, columns=["字段", "值"]),
                use_container_width=True,
                hide_index=True,
            )
            st.caption(
                "派生逻辑见 experiments/environment_engineering/derive.py（不进 FEATURE_COLUMNS）。"
            )

    def _show_fiber_engineering_bridge_section(self, user_inputs: dict) -> None:
        """开裂机理 Tab：纤维桥联能力解释卡 + KPI。"""
        feats = self._fiber_derive_row(user_inputs)
        st.markdown(section_title_html("纤维桥联能力解释"), unsafe_allow_html=True)
        summary = feats.get("fiber_engineering_summary") or {}
        kpis = [
            (
                "桥联能力",
                summary.get("expected_bridging_zh", "—").replace("预计桥联能力：", ""),
                f"FCI={summary.get('fiber_constraint_index_display', '—')}",
            ),
            (
                "界面协同",
                summary.get("interface_bond_zh", "—"),
                f"粘结={feats.get('fiber_interface_bond', '—')}",
            ),
            (
                "温度协调能力",
                summary.get("thermal_constraint_capacity_zh", "—").replace(
                    "温度变形协调能力：", ""
                ),
                f"E_f={summary.get('E_f_display', '—')}",
            ),
        ]
        st.markdown(fiber_engineering_kpi_html(kpis), unsafe_allow_html=True)
        cards = feats.get("bridge_explanation_zh") or []
        if cards:
            st.markdown(fiber_bridge_cards_html(cards), unsafe_allow_html=True)

    def _show_fiber_thermal_coordination_note(self, user_inputs: dict) -> None:
        note = self._fiber_derive_row(user_inputs).get("thermal_fiber_note_zh")
        if note:
            st.markdown(fiber_thermal_coordination_note_html(note), unsafe_allow_html=True)

    def _show_fiber_engineering_summary_advanced(self, user_inputs: dict) -> None:
        with st.expander("纤维工程体系 · fiber_engineering_summary", expanded=False):
            feats = self._fiber_derive_row(user_inputs)
            summary = feats.get("fiber_engineering_summary") or {}
            st.markdown(
                fiber_engineering_summary_panel_html(summary),
                unsafe_allow_html=True,
            )
            st.markdown("**派生量（解释层）**")
            rows = [
                ("fiber_diameter_mm", feats.get("fiber_diameter_mm")),
                ("fiber_constraint_index", feats.get("fiber_constraint_index")),
                ("aspect_ratio_norm", feats.get("aspect_ratio_norm")),
                ("bridge_capacity_score", feats.get("bridge_capacity_score")),
                ("thermal_coordination_score", feats.get("thermal_coordination_score")),
                ("interface_synergy_score", feats.get("interface_synergy_score")),
            ]
            st.dataframe(
                pd.DataFrame(rows, columns=["字段", "值"]),
                use_container_width=True,
                hide_index=True,
            )
            st.caption("派生逻辑见 experiments/fiber_engineering/derive.py（不进 FEATURE_COLUMNS）。")

    def show_fiber_influence_brief(self, user_inputs: dict, result: dict) -> None:
        st.markdown("##### 纤维与配合比（工程含义）")
        mat = user_inputs.get("fiber_material", "—")
        ft = user_inputs.get("fiber_type", "—")
        fc = user_inputs.get("fiber_content", "—")
        ar = user_inputs.get("aspect_ratio", "—")
        st.markdown(
            f"- 材质 **{mat}** · 外形 **{ft}** · 掺量 **{fc}%** · 长径比 **{ar}**  \n"
            "纤维主要通过影响抗裂能力、裂缝开展与残余变形，进入开裂风险预测模型的特征向量；"
            "详细数据驱动排序见下方「特征重要性」。"
        )
        mid = result.get("intermediate") or {}
        if mid.get("crack_density_source") == "fallback":
            st.caption(
                "提示：本次裂缝密度采用经验回退，纤维对密度的解释权重在机理图中可能偏弱。"
            )

    def show_tab_crack_mechanism(
        self,
        predictor,
        X_df: pd.DataFrame,
        result: dict,
        user_inputs: dict,
    ) -> None:
        """Tab3：开裂机理（工程化解释主路径）。"""
        st.markdown(tab_heading_html("开裂机理"), unsafe_allow_html=True)
        st.markdown(
            section_title_html("开裂形成路径解释"),
            unsafe_allow_html=True,
        )
        st.caption(
            "从材料—变形—约束—裂缝开展解释当前样本；"
            "**不替代** 主 Tab 中的裂缝宽度与开裂风险数值。"
        )
        self.show_mechanism_engineering_surface(
            predictor, X_df, result, user_inputs
        )
        with st.expander("高级分析（开发/研究）", expanded=False):
            st.caption(
                "以下为 PDP/ICE、全局特征重要性、训练 JSON 与温度诊断，"
                "默认仅供研发复核。"
            )
            self.show_mechanism_analysis(
                predictor, X_df, result, include_offline_json=True
            )
            with st.expander("温度应力 · 高级诊断", expanded=False):
                self.show_thermal_stress_rd_diagnostics(user_inputs)
            self._show_fiber_engineering_summary_advanced(user_inputs)
            self._show_environment_engineering_summary_advanced(user_inputs)

    @staticmethod
    def _input_range_status(
        value: float | None, lo: float, hi: float
    ) -> tuple[str, str]:
        if value is None or not math.isfinite(value):
            return "未输入", "muted"
        if lo <= value <= hi:
            return "正常", "low"
        span = max(hi - lo, 1e-9)
        margin = span * 0.2
        if value > hi:
            return ("偏高", "mid") if value <= hi + margin else ("偏高", "high")
        return ("偏低", "mid") if value >= lo - margin else ("偏低", "high")

    def _collect_input_experience_checks(self, user_inputs: dict) -> list[tuple[str, str, str]]:
        ranges = _cached_main_model_input_ranges()
        rows: list[tuple[str, str, str]] = []

        def add_num(key: str, label: str) -> None:
            tup = ranges.get(key)
            if not isinstance(tup, (list, tuple)) or len(tup) < 2:
                return
            lo, hi = float(tup[0]), float(tup[1])
            raw = user_inputs.get(key)
            try:
                v = float(raw)
                val: float | None = v if math.isfinite(v) else None
            except (TypeError, ValueError):
                val = None
            status, tone = self._input_range_status(val, lo, hi)
            rows.append((label, status, tone))

        add_num("w_b_ratio", "水胶比")
        add_num("fiber_content", "纤维掺量")
        add_num("aspect_ratio", "长径比")
        add_num("temperature", "环境温度")
        add_num("humidity", "环境湿度")

        merged = normalize_prediction_inputs(user_inputs)
        feats = derive_thermal_stress_features(series_for_thermal_derive(merged))
        if self._thermal_tsi_computable(feats):
            band = self._thermal_tsi_risk_band(float(feats["thermal_stress_index"]))
            if band == "高":
                rows.append(("温度路径", "指数偏高", "mid"))
            elif band == "中":
                rows.append(("温度路径", "一般", "low"))
            else:
                rows.append(("温度路径", "正常", "low"))
        else:
            rows.append(("温度路径", "未输入", "muted"))

        bounds = _cached_cube_strength_mpa_min_max_from_data_csv()
        grade = user_inputs.get("strength_grade")
        if bounds and isinstance(grade, str) and grade in STRENGTH_GRADE_TO_MPA:
            mpa = float(STRENGTH_GRADE_TO_MPA[grade])
            lo_b, hi_b = bounds
            status, tone = self._input_range_status(mpa, lo_b, hi_b)
            rows.append(("强度等级映射", status, tone))

        return rows

    @staticmethod
    def _assess_crack_model_stability() -> list[tuple[str, str, str]]:
        """由 training_metrics.json 映射工程稳定性文案（不展示 MAE/RMSE）。"""
        tm = load_training_metrics_json(MODELS_DIR)
        rows: list[tuple[str, str, str]] = []
        if not tm:
            return [
                ("裂缝宽度预测", "暂无离线报告", "muted"),
                ("开裂风险分类", "暂无离线报告", "muted"),
                ("裂缝密度预测", "暂无离线报告", "muted"),
            ]

        cw = tm.get("crack_width") if isinstance(tm.get("crack_width"), dict) else {}
        cd = tm.get("crack_density") if isinstance(tm.get("crack_density"), dict) else {}
        cr = tm.get("cracking_risk") if isinstance(tm.get("cracking_risk"), dict) else {}

        def reg_stability(block: dict) -> tuple[str, str]:
            r2 = block.get("test_r2")
            try:
                r2f = float(r2)
            except (TypeError, ValueError):
                return "待评估", "muted"
            if not math.isfinite(r2f):
                return "待评估", "muted"
            if r2f >= 0.4:
                return "相对稳定", "low"
            if r2f >= 0.1:
                return "一般", "mid"
            return "波动较大", "high"

        def cls_stability(block: dict) -> tuple[str, str]:
            acc = block.get("test_accuracy")
            try:
                af = float(acc)
            except (TypeError, ValueError):
                return "待评估", "muted"
            if not math.isfinite(af):
                return "待评估", "muted"
            if af >= 0.65:
                return "相对稳定", "low"
            if af >= 0.45:
                return "一般", "mid"
            return "波动较大", "high"

        rows.append(("裂缝宽度预测", *reg_stability(cw)))
        rows.append(("开裂风险分类", *cls_stability(cr)))
        rows.append(("裂缝密度预测", *reg_stability(cd)))
        return rows

    def _build_governance_engineering_bullets(self) -> list[str]:
        gov_path = PROJECT_ROOT / "outputs" / "crack_governance" / "crack_training_governance.json"
        rep = self._try_load_json_dict(gov_path)
        if not rep:
            return [
                "治理列尚未生成诊断报告，协议分层信息未接入展示",
                "当前结论主要依据通用训练表与离线指标，请谨慎用于主结论",
            ]
        cols = rep.get("column_presence") or {}
        if all(cols.get(c) for c in (
            "source_group",
            "data_tier",
            "crack_width_definition_id",
        )):
            bullets = [
                "已接入试验协议与来源治理列（侧车字段）",
                "已区分 A/B/C 数据层级与暂缓（hold）机制",
            ]
        else:
            bullets = ["部分治理列尚未写入训练表，协议分层能力不完整"]
        tc = rep.get("tier_ABC_hold_counts") or {}
        hold = int(tc.get("hold_pending", 0) or 0)
        tier_a = int(tc.get("tier_A_candidate", 0) or 0)
        if tier_a > 0:
            bullets.append("训练库中已有 A 类高可信样本可用于主结论对照")
        elif hold > 0:
            bullets.append(
                "当前训练数据多数仍处于暂缓/未分级状态，尚难进入高可信协议集"
            )
        ga = rep.get("group_audit") or {}
        if int(ga.get("n_rows_with_nonempty_source_group", 0) or 0) > 0:
            bullets.append("已记录来源分组，便于分组评估与泄漏排查")
        else:
            bullets.append("来源分组字段尚未填写，分组可信度评估能力受限")
        bullets.append("当前单次侧栏输入未绑定文献/试验协议编号，按通用模型推断")
        return bullets[:5]

    def _build_trust_conclusion(
        self,
        user_inputs: dict,
        result: dict,
        checks: list[tuple[str, str, str]],
        stability: list[tuple[str, str, str]],
    ) -> tuple[str, str, str, list[str], list[str]]:
        """返回 level_label, tone, intro, positives, caveats。"""
        assessment = assess_trust(
            checks=checks,
            stability=stability,
            result=result,
            user_inputs=user_inputs,
        )
        return (
            assessment.level_label,
            assessment.tone,
            assessment.intro,
            assessment.positives,
            assessment.caveats,
        )

    def _show_trust_advanced_offline_metrics(self) -> None:
        """研发向：MAE/RMSE/OOF/residual/JSON 路径。"""
        st.caption(CRACK_OFFLINE_METRICS_DISCLAIMER)
        crack_tm_path = PROJECT_ROOT / "outputs" / "training_metrics.json"
        lab_json_path = (
            PROJECT_ROOT / "outputs" / "lab_strength" / "lab_strength_residual_report.json"
        )
        tm_rep = self._try_load_json_dict(crack_tm_path)
        st.markdown("##### 主开裂模型（历史 hold-out 测试集）")
        st.markdown(self._crack_main_offline_markdown(tm_rep))
        with st.expander("training_metrics.json 原文", expanded=False):
            if tm_rep:
                st.json(tm_rep)
            else:
                st.caption(f"未找到 `{crack_tm_path}`")

        lab_rep = self._try_load_json_dict(lab_json_path)
        st.markdown("---")
        st.markdown("##### 力学强度管线（历史训练）")
        st.markdown(self._lab_strength_offline_markdown(lab_rep))
        with st.expander("lab_strength_residual_report.json 原文", expanded=False):
            if lab_rep:
                st.json(lab_rep)
            else:
                st.caption(f"未找到 `{lab_json_path}`")

        self._render_lab_strength_offline_metrics_body()

        gov_path = PROJECT_ROOT / "outputs" / "crack_governance" / "crack_training_governance.json"
        with st.expander("crack_training_governance.json", expanded=False):
            g = self._try_load_json_dict(gov_path)
            if g:
                st.json(g)
            else:
                st.caption(f"未找到 `{gov_path}`")
        st.caption(f"路径：`{crack_tm_path}` · `{lab_json_path}` · `{gov_path}`")

    def show_tab_data_trust(self, user_inputs: dict, result: dict) -> None:
        """Tab4：模型可信度（工程评估主路径）。"""
        st.markdown(tab_heading_html("模型可信度"), unsafe_allow_html=True)
        st.markdown(
            section_title_html("预测结果可信度评估"),
            unsafe_allow_html=True,
        )
        st.caption(
            "回答「当前结果靠不靠谱」：结合输入经验范围、分路径公式/残差方法论、"
            "世界数据库治理状态与历史 hold-out 稳定性；不展示单次试验误差。"
        )

        checks = self._collect_input_experience_checks(user_inputs)
        stability = self._assess_crack_model_stability()
        assessment = assess_trust(
            checks=checks,
            stability=stability,
            result=result,
            user_inputs=user_inputs,
        )
        level, tone, intro, positives, caveats = (
            assessment.level_label,
            assessment.tone,
            assessment.intro,
            assessment.positives,
            assessment.caveats,
        )

        st.markdown(
            trust_conclusion_banner_html(level, tone, intro, positives, caveats),
            unsafe_allow_html=True,
        )
        st.markdown(
            trust_score_chip_html(assessment.trust_score),
            unsafe_allow_html=True,
        )
        pipe_rows = [
            {
                "task": p.task,
                "method": p.method,
                "evidence": p.evidence,
                "stability": p.stability,
                "tone": p.tone,
                "note": p.note,
            }
            for p in assessment.pipelines
        ]
        st.markdown(
            trust_pipeline_methodology_html(pipe_rows),
            unsafe_allow_html=True,
        )

        st.markdown(
            input_range_check_panel_html("输入经验范围检查", checks),
            unsafe_allow_html=True,
        )
        st.caption(
            "区间来自训练数据观测 min/max（部分参数回退常用工程带）；"
            "仅提示、不拦截预测。"
        )

        st.markdown(
            stability_status_panel_html(stability),
            unsafe_allow_html=True,
        )
        st.caption(
            "上表「历史模型稳定性」标签由历史 hold-out 指标映射，"
            "不代表当前输入实时误差；权重与指标为弱同源。"
        )
        with st.expander("查看历史 hold-out 指标（开发/研究）", expanded=False):
            st.caption(CRACK_OFFLINE_METRICS_DISCLAIMER)
            tm_rep = self._try_load_json_dict(
                PROJECT_ROOT / "outputs" / "training_metrics.json"
            )
            st.markdown(self._crack_main_offline_markdown(tm_rep))

        st.markdown(
            trust_notice_card_html(
                "数据治理状态",
                self._build_governance_engineering_bullets(),
            ),
            unsafe_allow_html=True,
        )

        st.markdown(
            standards_compact_html(
                [
                    "GB/T 50081-2019《普通混凝土力学性能试验方法标准》",
                    "GB 50010-2010（2015 年版）《混凝土结构设计规范》",
                    "GB/T 50082-2009《普通混凝土长期性能和耐久性能试验方法标准》",
                    "《大体积混凝土施工技术规程》",
                    "GB 50164-2011《混凝土质量控制标准》",
                    "JGJ/T 221-2010《纤维混凝土应用技术规程》",
                ],
                "本系统用于工程辅助评估，不能替代正式试验与结构验算。",
            ),
            unsafe_allow_html=True,
        )

        with st.expander("高级历史 hold-out 指标（开发/研究）", expanded=False):
            self._show_trust_advanced_offline_metrics()

    def show_heuristic_time_metrics_rd(self, result: dict) -> None:
        preds = result.get("predictions", {}) or {}
        td = preds.get("time_dimension") or {}
        if not td:
            return
        st.markdown("##### 时间维度（启发式，非试验测定）")
        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric(
                "开裂时间（启发式）",
                self._fmt_float(td.get("cracking_time_hours"), 1, " h"),
            )
        with c2:
            st.metric(
                "临界龄期（启发式）",
                self._fmt_float(td.get("critical_age_days"), 2, " d"),
            )
        with c3:
            st.metric(
                "安全窗口（启发式）",
                self._fmt_float(td.get("safety_window_hours"), 1, " h"),
            )
        if td.get("note"):
            st.caption(td["note"])

    def _show_thermal_advanced_diagnostics(
        self,
        merged: dict,
        display: dict,
        feats: dict,
    ) -> None:
        """missing_flag、derive 路径、JSON 与无量纲指数（研发向）。"""
        st.caption(
            "以下为派生链路与 missing_flag，供复核与对接 CSV/脚本；"
            "默认用户无需展开。"
        )
        flag_rows = [
            ("delta_T_eff", feats.get("delta_T_eff_missing_flag")),
            ("cooling_rate", feats.get("cooling_rate_missing_flag")),
            ("thermal_gradient_index", feats.get("thermal_gradient_index_missing_flag")),
            ("E_norm", feats.get("E_norm_missing_flag")),
            ("alpha_norm", feats.get("alpha_norm_missing_flag")),
            ("restraint_factor_R", feats.get("restraint_factor_R_missing_flag")),
            ("thermal_stress_index", feats.get("thermal_stress_index_missing_flag")),
            ("thermal_crack_risk_index", feats.get("thermal_crack_risk_index_missing_flag")),
            ("engineering_chain", display.get("engineering_chain_missing_flag")),
        ]
        st.markdown("**missing_flag（1=缺失/不可算）**")
        st.dataframe(
            pd.DataFrame(flag_rows, columns=["派生项", "missing_flag"]),
            use_container_width=True,
            hide_index=True,
            height=320,
        )
        src = feats.get("delta_T_eff_source")
        st.markdown(
            f"**ΔT 派生路径（derive）：** `{src}` → "
            f"{self._thermal_delta_t_source_zh(src)}"
        )
        st.markdown("**工程公式层数值**")
        eng_rows = [
            ("epsilon_T", display.get("epsilon_T")),
            ("sigma_T_explain", display.get("sigma_T_explain")),
            ("f_t_proxy_MPa", display.get("f_t_proxy_MPa")),
            ("eta", display.get("eta")),
            ("eta_risk_band", display.get("eta_risk_band")),
            ("E_MPa", display.get("E_MPa")),
            ("E_source", display.get("E_source")),
            ("f_t_proxy_source", display.get("f_t_proxy_source")),
            ("f_t_proxy_source_zh", display.get("f_t_proxy_source_zh")),
        ]
        st.dataframe(
            pd.DataFrame(eng_rows, columns=["变量", "值"]),
            use_container_width=True,
            hide_index=True,
        )
        opt_ctx = display.get("optional_context")
        if opt_ctx and self._thermal_optional_context_has_data(opt_ctx):
            st.markdown("**扩展可选输入（optional_context）**")
            self._show_thermal_optional_context_panel(opt_ctx)
        st.markdown("**Phase 1 派生量（调试）**")
        debug_rows = [
            ("thermal_stress_index", feats.get("thermal_stress_index")),
            ("delta_T_eff", feats.get("delta_T_eff")),
            ("restraint_factor_R", feats.get("restraint_factor_R")),
            ("E_norm", feats.get("E_norm")),
            ("alpha_norm", feats.get("alpha_norm")),
            ("thermal_crack_risk_index", feats.get("thermal_crack_risk_index")),
        ]
        st.dataframe(
            pd.DataFrame(debug_rows, columns=["变量名", "值"]),
            use_container_width=True,
            hide_index=True,
        )
        with st.expander("JSON 原文（feats + 工程链）", expanded=False):
            st.json(
                {
                    "derive_thermal_stress_features": {
                        k: feats[k]
                        for k in sorted(feats.keys())
                        if k != "feats"
                    },
                    "engineering_display": {
                        k: display[k]
                        for k in sorted(display.keys())
                        if k not in ("feats", "f_t_proxy_detail")
                    },
                }
            )
        if display.get("f_t_proxy_detail"):
            with st.expander("f_t 映射明细", expanded=False):
                st.json(display["f_t_proxy_detail"])
        st.caption("派生逻辑见 experiments/thermal_stress/derive.py · engineering_chain.py")

    @staticmethod
    def _thermal_optional_context_has_data(opt_ctx: dict) -> bool:
        if not opt_ctx:
            return False
        for k, v in opt_ctx.items():
            if k.endswith("_missing_flag") or k.endswith("_zh") or k.endswith("_note_zh"):
                continue
            if v is not None and str(v).strip().lower() not in ("", "nan", "none"):
                return True
        return False

    def _show_thermal_optional_context_panel(self, opt_ctx: dict) -> None:
        """展示扩展可选输入（温度路径、试验约束、裂缝观测）。"""
        if opt_ctx.get("temperature_path_summary_zh"):
            st.markdown(f"**温度路径：** {opt_ctx['temperature_path_summary_zh']}")
        if opt_ctx.get("restraint_test_display_zh"):
            st.caption(opt_ctx.get("restraint_test_note_zh", ""))
            st.markdown(f"**试验约束：** {opt_ctx['restraint_test_display_zh']}")
        elif opt_ctx.get("restraint_test_note_zh"):
            st.caption(opt_ctx["restraint_test_note_zh"])

        crack_rows = []
        if opt_ctx.get("thermal_crack_observed") is not None:
            crack_rows.append(
                ("温度裂缝", "是" if opt_ctx["thermal_crack_observed"] else "否")
            )
        if opt_ctx.get("thermal_crack_time_h") is not None:
            crack_rows.append(("裂缝时间 (h)", opt_ctx["thermal_crack_time_h"]))
        if opt_ctx.get("thermal_crack_width_mm") is not None:
            crack_rows.append(("裂缝宽度 (mm)", opt_ctx["thermal_crack_width_mm"]))
        if opt_ctx.get("apparent_crack_filtered") is not None:
            crack_rows.append(
                (
                    "表观裂缝已过滤",
                    "是" if opt_ctx["apparent_crack_filtered"] else "否",
                )
            )
        if crack_rows:
            if opt_ctx.get("crack_validation_note_zh"):
                st.caption(opt_ctx["crack_validation_note_zh"])
            st.dataframe(
                pd.DataFrame(crack_rows, columns=["项", "值"]),
                use_container_width=True,
                hide_index=True,
            )

        st.markdown("**可选字段 missing_flag（1=未提供）**")
        flag_keys = [k for k in sorted(opt_ctx.keys()) if k.endswith("_missing_flag")]
        if flag_keys:
            st.dataframe(
                pd.DataFrame(
                    [(k.replace("_missing_flag", ""), opt_ctx[k]) for k in flag_keys],
                    columns=["字段", "missing_flag"],
                ),
                use_container_width=True,
                hide_index=True,
                height=min(280, 35 + 28 * len(flag_keys)),
            )

    def show_thermal_stress_rd_diagnostics(self, user_inputs: dict) -> None:
        """研发工作台：温度应力派生链诊断（展示层，不改 derive）。"""
        merged = normalize_prediction_inputs(user_inputs)
        row = series_for_thermal_derive(merged)
        display = derive_thermal_engineering_display(row)
        feats = display.get("feats") or derive_thermal_stress_features(row)
        self._show_thermal_advanced_diagnostics(merged, display, feats)

    def _build_rd_status_overview_cards(self) -> list[dict]:
        tm = load_training_metrics_json(MODELS_DIR)
        model_val = "Stable" if tm else "待加载"
        model_tone = "blue" if tm else "yellow"

        gov_path = PROJECT_ROOT / "outputs" / "crack_governance" / "crack_training_governance.json"
        rep = self._try_load_json_dict(gov_path)
        tc = rep.get("tier_ABC_hold_counts") or {} if rep else {}
        hold = int(tc.get("hold_pending", 0) or 0)
        tier_a = int(tc.get("tier_A_candidate", 0) or 0)
        gov_val = "Pending" if hold > 0 else "就绪"
        gov_tone = "yellow" if hold > 0 else "green"
        tier_tone = "green" if tier_a > 0 else "red"

        return [
            {"label": "模型状态", "value": model_val, "tone": model_tone},
            {"label": "数据治理", "value": gov_val, "tone": gov_tone},
            {"label": "A 类协议集", "value": str(tier_a), "tone": tier_tone},
            {"label": "温度应力模块", "value": "Phase 1", "tone": "cyan"},
        ]

    def _governance_tier_counts(self) -> tuple[int, int, int, int]:
        gov_path = PROJECT_ROOT / "outputs" / "crack_governance" / "crack_training_governance.json"
        rep = self._try_load_json_dict(gov_path)
        if not rep:
            return 0, 0, 0, 0
        tc = rep.get("tier_ABC_hold_counts") or {}
        return (
            int(tc.get("tier_A_candidate", 0) or 0),
            int(tc.get("tier_B", 0) or 0),
            int(tc.get("hold_pending", 0) or 0),
            int(tc.get("tier_C_or_illegal_tier", 0) or 0),
        )

    def _render_rd_module_model_diagnosis(
        self,
        result: dict,
        user_inputs: dict,
    ) -> None:
        st.markdown(
            research_module_header_html(
                "模型诊断",
                "模型训练与预测行为分析",
                badge="Research",
            ),
            unsafe_allow_html=True,
        )
        st.markdown(
            '<p class="sfc-rd-note"><strong>模型行为与输入输出诊断</strong> · '
            "用于分析模型当前预测逻辑与派生链，不代表正式工程验算。</p>",
            unsafe_allow_html=True,
        )
        with st.expander("训练指标摘要", expanded=True):
            tm = load_training_metrics_json(MODELS_DIR)
            if tm:
                st.json(tm)
            else:
                st.caption("未找到训练指标文件。")
            st.caption("源文件：models/training_metrics.json")
            cv = load_cv_metrics_json(MODELS_DIR)
            if cv:
                st.markdown("**交叉验证摘要**")
                st.json(cv)
                st.caption("源文件：models/cv_metrics.json")

        with st.expander("主开裂 · 输入输出明细", expanded=False):
            self.show_detail_table(result, user_inputs)

        with st.expander("温度应力 · 高级诊断", expanded=False):
            st.caption(
                "温度应力为解释指数，不代表 MPa 级应力；工程结论见「③ 开裂机理」。"
            )
            self.show_thermal_stress_rd_diagnostics(user_inputs)
            st.caption("派生逻辑见 experiments/thermal_stress/derive.py（只读）。")

    def _render_rd_module_data_governance(self) -> None:
        tier_a, tier_b, hold, tier_c = self._governance_tier_counts()
        st.markdown(
            research_module_header_html(
                "数据治理",
                "训练数据协议分层与样本质量",
                badge="Governance",
            ),
            unsafe_allow_html=True,
        )
        with st.expander("数据治理 · 协议与样本", expanded=False):
            st.markdown(
                governance_tier_strip_html(tier_a, tier_b, hold, tier_c),
                unsafe_allow_html=True,
            )
            st.markdown(
                '<p class="sfc-rd-note">当前数据仍以未治理样本为主，'
                "不建议直接用于正式重训。</p>",
                unsafe_allow_html=True,
            )
            with st.expander("数据治理状态 · 完整报告", expanded=False):
                gov_path = (
                    PROJECT_ROOT
                    / "outputs"
                    / "crack_governance"
                    / "crack_training_governance.json"
                )
                g = self._try_load_json_dict(gov_path)
                if g:
                    st.json(g)
                else:
                    st.caption("未找到治理诊断报告。")
                st.caption("源文件：outputs/crack_governance/crack_training_governance.json")

            drop_path = (
                PROJECT_ROOT / "outputs" / "lab_strength" / "lab_strength_dropped_rows.json"
            )
            rep_path = (
                PROJECT_ROOT
                / "outputs"
                / "lab_strength"
                / "lab_strength_residual_report.json"
            )
            with st.expander("协议不完整样本", expanded=False):
                dropped = self._try_load_json_dict(drop_path)
                if dropped:
                    st.json(dropped)
                else:
                    st.caption("暂无记录。")
                st.caption("源文件：outputs/lab_strength/lab_strength_dropped_rows.json")

            with st.expander("人工复核样本", expanded=False):
                rep = self._try_load_json_dict(rep_path)
                if rep:
                    n = rep.get("manual_review_rows_count", "—")
                    st.markdown(f"**人工复核行数（报告统计）：** {n}")
                    if rep.get("manual_review_samples"):
                        st.json(rep.get("manual_review_samples"))
                else:
                    st.caption("暂无力学强度报告。")
                st.caption("源文件：outputs/lab_strength/lab_strength_residual_report.json")

    def _render_rd_module_research_experiments(self, result: dict) -> None:
        st.markdown(
            research_module_header_html(
                "研究实验",
                "实验性算法与启发式扩展（非主结论）",
                badge="Experimental",
                badge_class="sfc-rd-badge-exp",
            ),
            unsafe_allow_html=True,
        )
        with st.expander("研究实验 · 实验性模块", expanded=False):
            st.caption("实验性模块 / Experimental · 结果不进入主工程结论。")
            with st.expander("生成式扩增实验（GAN）", expanded=False):
                st.markdown(
                    "**状态：已冻结。** 不纳入主训练、不用于提升主模型指标。"
                    "历史产物目录：experiments/gan/"
                )

            with st.expander("启发式时间指标与应力-强度比", expanded=False):
                self.show_heuristic_time_metrics_rd(result)
                sd = (result.get("predictions") or {}).get("state_dimension") or {}
                if sd.get("stress_strength_ratio") is not None:
                    st.metric(
                        "应力/抗拉强度比（启发式）",
                        self._fmt_float(sd.get("stress_strength_ratio"), 3),
                    )
                if sd.get("stress_strength_note_cn"):
                    st.caption(sd["stress_strength_note_cn"])

            with st.expander("优化建议与敏感性（规划）", expanded=False):
                self.show_optimization_suggestions(result)
                self.show_sensitivity(result)

    def _render_rd_module_export(
        self,
        result: dict,
        user_inputs: dict,
        predictor,
        X_df: pd.DataFrame,
    ) -> None:
        st.markdown(
            research_module_header_html(
                "导出与记录",
                "报告导出与结果留档",
                badge="Export",
            ),
            unsafe_allow_html=True,
        )
        with st.expander("导出与记录", expanded=False):
            with st.expander("Word 报告导出", expanded=False):
                self.show_report_export(result, user_inputs, predictor, X_df)

    def show_tab_rd_diagnostics(
        self,
        result: dict,
        user_inputs: dict,
        predictor,
        X_df: pd.DataFrame,
    ) -> None:
        """Tab5：高级分析 · 研发工作台。"""
        st.markdown(tab_heading_html("高级分析"), unsafe_allow_html=True)
        st.markdown(
            section_title_html("研发工作台 · Research Workspace"),
            unsafe_allow_html=True,
        )
        st.caption(
            "模型研发与诊断层，供研究人员排查训练、治理与实验模块；"
            "**不影响**侧栏实时预测计算。"
        )
        st.markdown(
            research_status_overview_html(
                "研发状态总览",
                self._build_rd_status_overview_cards(),
            ),
            unsafe_allow_html=True,
        )
        self._render_rd_module_model_diagnosis(result, user_inputs)
        render_model_performance_evaluation(plot_fn=self._plotly)
        self._show_fiber_engineering_summary_advanced(user_inputs)
        self._show_environment_engineering_summary_advanced(user_inputs)
        self._render_rd_module_data_governance()
        self._render_rd_module_research_experiments(result)
        self._render_rd_module_export(result, user_inputs, predictor, X_df)

    def show_result_cards(self, result: dict) -> None:
        """兼容旧入口：等同综合预测核心 + 研发向时间维度（新 UI 请用 show_tab_*）。"""
        self.show_tab_comprehensive_prediction(result, {})
        preds = result.get("predictions", {}) or {}
        td = preds.get("time_dimension") or {}
        if td:
            self.show_heuristic_time_metrics_rd(result)

    def show_risk_gauge(self, result: dict) -> None:
        st.markdown(
            '<div class="sfc-gauge-compact"><p class="sfc-zone-tight">'
            + section_title_html("开裂风险概率监测")
            + "</p></div>",
            unsafe_allow_html=True,
        )
        preds = result.get("predictions", {}) or {}
        sd = preds.get("state_dimension") or {}
        try:
            p = float(sd.get("risk_probability", preds.get("risk_confidence", 0.0)))
        except (TypeError, ValueError):
            p = 0.0
        if not math.isfinite(p):
            p = 0.0
        p = min(max(p, 0.0), 1.0)
        pct = p * 100.0
        if pct < 30:
            bar_c = "#94a3b8"
        elif pct < 70:
            bar_c = "#64748b"
        else:
            bar_c = "#b91c1c"

        try:
            fig = go.Figure(
                go.Indicator(
                    mode="gauge+number",
                    value=pct,
                    number={
                        "font": {"size": 32, "color": COLORS["text_muted"]},
                        "suffix": " %",
                        "valueformat": ".1f",
                    },
                    title={
                        "text": "开裂风险概率 P",
                        "font": {"size": 13, "color": COLORS["text_muted"]},
                    },
                    gauge={
                        "shape": "angular",
                        "axis": {
                            "range": [0, 100],
                            "tickcolor": COLORS["border"],
                            "tickfont": {"size": 11, "color": COLORS["text_muted"]},
                            "tickwidth": 1,
                        },
                        "bar": {"color": bar_c, "thickness": 0.22},
                        "bgcolor": COLORS["plot"],
                        "borderwidth": 0,
                        "steps": [
                            {"range": [0, 30], "color": "rgba(148, 163, 184, 0.18)"},
                            {"range": [30, 70], "color": "rgba(100, 116, 139, 0.22)"},
                            {"range": [70, 100], "color": "rgba(185, 28, 28, 0.22)"},
                        ],
                        "threshold": {
                            "line": {"color": bar_c, "width": 3},
                            "thickness": 0.78,
                            "value": pct,
                        },
                    },
                )
            )
            fig = style_gauge_indicator(fig, height=210)
            self._plotly(fig)
            st.caption(
                "色带：0–30% 低风险 · 30–70% 中风险 · 70–100% 高风险；指针颜色与当前所在预警带一致。"
            )
        except Exception:
            st.metric("开裂风险概率 P", f"{pct:.1f} %")
            st.progress(min(max(p, 0.0), 1.0))

    @staticmethod
    def _thermal_tsi_computable(feats: dict) -> bool:
        tsi_mf = feats.get("thermal_stress_index_missing_flag")
        tsi = feats.get("thermal_stress_index")
        if tsi_mf == 1:
            return False
        try:
            v = float(tsi)
        except (TypeError, ValueError):
            return False
        return math.isfinite(v) and v >= 0.0

    @staticmethod
    def _thermal_tsi_risk_band(tsi: float) -> str:
        if tsi >= 0.35:
            return "高"
        if tsi >= 0.18:
            return "中"
        return "低"

    @staticmethod
    def _thermal_missing_input_labels(user_inputs: dict, feats: dict) -> list[str]:
        """工程语言：列出阻断完整评估的缺失项。"""
        missing: list[str] = []
        if feats.get("delta_T_eff_missing_flag") == 1:
            missing.append("温差 ΔT（侧栏「内外/芯表温差」或等价温差路径）")
        if feats.get("restraint_factor_R_missing_flag") == 1:
            missing.append("约束等级（侧栏请选择 low / medium / high，勿留 unknown）")
        if feats.get("alpha_norm_missing_flag") == 1:
            missing.append("线膨胀系数 α（×10⁻⁶/℃）")
        if feats.get("E_norm_missing_flag") == 1:
            missing.append("材料刚度 E（弹性模量或可用的强度等级代理）")
        return missing

    @staticmethod
    def _thermal_delta_t_source_zh(src: str | None) -> str:
        mapping = {
            "delta_T_inner_outer": "内外/芯表温差",
            "delta_T_user": "用户给定控制温差",
            "T_peak_minus_T_reference": "峰值温度 − 参考温度",
            "missing": "未提供温差路径",
        }
        if not src:
            return "—"
        return mapping.get(str(src), str(src))

    @staticmethod
    def _thermal_restraint_label_zh(user_inputs: dict) -> str:
        rl = str(user_inputs.get("restraint_level", "unknown")).strip().lower()
        return {
            "low": "低约束",
            "medium": "中约束",
            "high": "高约束",
            "unknown": "未指定",
        }.get(rl, rl)

    def _render_thermal_chain_card(
        self,
        *,
        title: str,
        status: str,
        value_line: str,
        help_line: str,
    ) -> None:
        st.markdown(f"**{title}**")
        st.markdown(f"{status}　{value_line}")
        st.caption(help_line)

    def show_thermal_stress_phase1(
        self, user_inputs: dict, *, embedded: bool = False
    ) -> None:
        """温度应力解释（Phase 1）：工程解释链展示，不替代主模型 MPa/mm 输出。"""
        def _body() -> None:
            merged = normalize_prediction_inputs(user_inputs)
            row = series_for_thermal_derive(merged)
            feats = derive_thermal_stress_features(row)
            computable = self._thermal_tsi_computable(feats)
            missing_labels = self._thermal_missing_input_labels(merged, feats)

            # —— 1. 顶部工程结论 ——
            st.markdown("#### 工程结论")
            if not computable:
                st.warning("**当前无法形成完整温度应力评估**")
                if missing_labels:
                    st.markdown("**缺少：**")
                    for item in missing_labels:
                        st.markdown(f"- {item}")
                st.markdown(
                    "**建议：** 在左侧「温度应力解释（Phase 1，可选）」中补全上述项后，"
                    "本页将自动重新计算。"
                )
            else:
                tsi_f = float(feats["thermal_stress_index"])
                band = self._thermal_tsi_risk_band(tsi_f)
                st.success(f"**温度应力指数：{band}**（数值约 {tsi_f:.2f}）")
                st.markdown(thermal_stress_explain_sentence_zh(feats))

            st.markdown(
                "温度应力会放大约束与材料变形不匹配带来的开裂倾向；"
                "本模块用于**辅助理解**温度因素，**不替代**下方主模型给出的裂缝宽度与开裂风险。"
            )

            st.divider()

            # —— 2. 指数（弱化变量名）——
            st.markdown("#### 温度应力指数（Phase 1）")
            if not computable:
                st.markdown("**当前：不可计算**")
            else:
                tsi_f = float(feats["thermal_stress_index"])
                st.markdown(
                    f"**当前：{tsi_f:.2f}（{self._thermal_tsi_risk_band(tsi_f)}风险）**"
                )
            st.caption(
                "该指标仅用于解释温度变化与开裂风险的关系，不代表真实 MPa 级温度应力；"
                "不替代 crack_width / cracking_risk 等主模型输出。"
            )

            st.divider()

            # —— 3. 工程解释链（流程 + 分量）——
            st.markdown("#### 工程解释链")
            dt_ok = feats.get("delta_T_eff_missing_flag") == 0
            r_ok = feats.get("restraint_factor_R_missing_flag") == 0
            e_ok = feats.get("E_norm_missing_flag") == 0
            a_ok = feats.get("alpha_norm_missing_flag") == 0
            gradient_ok = dt_ok and (e_ok or a_ok)
            st.markdown(
                thermal_engineering_chain_html(
                    dt_ok=dt_ok,
                    gradient_ok=gradient_ok,
                    restraint_ok=r_ok,
                ),
                unsafe_allow_html=True,
            )
            c1, c2, c3, c4 = st.columns(4)

            dte = feats.get("delta_T_eff")
            dt_ok = feats.get("delta_T_eff_missing_flag") == 0
            dt_val = (
                f"{float(dte):.1f} °C"
                if dt_ok and self._is_good_num(dte)
                else "—"
            )
            with c1:
                self._render_thermal_chain_card(
                    title="温差 ΔT",
                    status="已输入" if dt_ok else "未输入",
                    value_line=dt_val,
                    help_line="构件内外或控制工况下的有效温差，是温度驱动的主项。",
                )

            r_ok = feats.get("restraint_factor_R_missing_flag") == 0
            Rv = feats.get("restraint_factor_R")
            r_val = (
                f"{self._thermal_restraint_label_zh(merged)}"
                + (
                    f"（系数约 {float(Rv):.2f}）"
                    if r_ok and self._is_good_num(Rv) and float(Rv) >= 0.0
                    else ""
                )
                if r_ok
                else "—"
            )
            with c2:
                self._render_thermal_chain_card(
                    title="约束条件",
                    status="已输入" if r_ok else "未输入",
                    value_line=r_val,
                    help_line="边界、厚度、滑移层与岩体/钢筋约束越强，温度应力越难释放。",
                )

            e_ok = feats.get("E_norm_missing_flag") == 0
            e_user = merged.get("elastic_modulus_E_user")
            e_note = (
                "侧栏已填弹性模量"
                if e_user is not None and self._is_good_num(e_user)
                else "由强度等级估算刚度"
            )
            En = feats.get("E_norm")
            e_val = (
                f"{e_note}（归一化 {float(En):.2f}）"
                if e_ok and self._is_good_num(En)
                else "—"
            )
            with c3:
                self._render_thermal_chain_card(
                    title="材料刚度 E",
                    status="已输入" if e_ok else "未输入",
                    value_line=e_val,
                    help_line="弹性模量越大，同等温差下越不易协调变形，应力解释项越高。",
                )

            a_ok = feats.get("alpha_norm_missing_flag") == 0
            alpha = merged.get("thermal_expansion_alpha")
            a_val = "—"
            if a_ok and alpha is not None and self._is_good_num(alpha):
                a_val = f"{float(alpha) * 1e6:.1f}×10⁻⁶/℃"
            with c4:
                self._render_thermal_chain_card(
                    title="热膨胀能力 α",
                    status="已输入" if a_ok else "未输入",
                    value_line=a_val,
                    help_line="线膨胀系数越大，温度变化时变形需求越大。",
                )

            st.divider()

            # —— 4. 高级诊断（默认折叠）——
            with st.expander("高级诊断（开发/研究）", expanded=False):
                st.caption("以下为派生链路与 missing_flag，供复核与对接 CSV/脚本。")
                flag_rows = [
                    ("delta_T_eff", feats.get("delta_T_eff_missing_flag")),
                    ("cooling_rate", feats.get("cooling_rate_missing_flag")),
                    ("thermal_gradient_index", feats.get(
                        "thermal_gradient_index_missing_flag"
                    )),
                    ("E_norm", feats.get("E_norm_missing_flag")),
                    ("alpha_norm", feats.get("alpha_norm_missing_flag")),
                    ("restraint_factor_R", feats.get("restraint_factor_R_missing_flag")),
                    ("thermal_stress_index", feats.get("thermal_stress_index_missing_flag")),
                    ("thermal_crack_risk_index", feats.get(
                        "thermal_crack_risk_index_missing_flag"
                    )),
                ]
                st.markdown("**missing_flag（1=缺失/不可算）**")
                st.dataframe(
                    pd.DataFrame(flag_rows, columns=["派生项", "missing_flag"]),
                    use_container_width=True,
                    hide_index=True,
                    height=280,
                )
                src = feats.get("delta_T_eff_source")
                st.markdown(
                    f"**ΔT 派生路径（derive）：** "
                    f"`{src}` → {self._thermal_delta_t_source_zh(src)}"
                )
                st.markdown("**派生量数值（调试）**")
                debug_rows = [
                    ("thermal_stress_index", feats.get("thermal_stress_index")),
                    ("delta_T_eff", feats.get("delta_T_eff")),
                    ("restraint_factor_R", feats.get("restraint_factor_R")),
                    ("E_norm", feats.get("E_norm")),
                    ("alpha_norm", feats.get("alpha_norm")),
                    ("cooling_rate", feats.get("cooling_rate")),
                    ("thermal_gradient_index", feats.get("thermal_gradient_index")),
                    ("thermal_crack_risk_index", feats.get("thermal_crack_risk_index")),
                ]
                st.dataframe(
                    pd.DataFrame(debug_rows, columns=["变量名", "值"]),
                    use_container_width=True,
                    hide_index=True,
                )

        if embedded:
            st.markdown(
                section_title_html("温度应力解释指数（Phase 1）"),
                unsafe_allow_html=True,
            )
            st.caption(
                "温度应力解释指数，不代表真实 MPa 应力；不替代 crack_width / cracking_risk。"
            )
            _body()
        else:
            with st.expander("温度应力解释（Phase 1）", expanded=False):
                _body()

    def show_detail_table(self, result: dict, user_inputs: dict) -> None:
        preds = result.get("predictions", {}) or {}
        flat_out: dict = {}
        td = preds.get("time_dimension")
        sd = preds.get("state_dimension")
        rd = preds.get("risk_dimension")
        if isinstance(td, dict):
            for k, v in td.items():
                if k != "note":
                    flat_out[f"时间.{k}"] = v
        if isinstance(sd, dict):
            for k, v in sd.items():
                if k != "note":
                    flat_out[f"状态.{k}"] = v
        if isinstance(rd, dict):
            for k, v in rd.items():
                if k != "bands":
                    flat_out[f"风险.{k}"] = v
        for k, v in preds.items():
            if k in ("time_dimension", "state_dimension", "risk_dimension"):
                continue
            flat_out[f"兼容.{k}"] = v

        def _cell(v):
            if isinstance(v, float) and not math.isfinite(v):
                return "—"
            if isinstance(v, dict):
                return str({k: _cell(x) for k, x in v.items()})
            return v

        flat_display = {k: _cell(v) for k, v in flat_out.items()}
        self._sync_prediction_history(user_inputs, flat_display)

        df_in = pd.DataFrame([{f"输入.{k}": v for k, v in user_inputs.items()}])
        df_pred = pd.DataFrame([flat_display])
        df_combo = pd.concat([df_in, df_pred], axis=1)

        st.markdown(
            section_title_html("输入参数与预测结果概览"),
            unsafe_allow_html=True,
        )
        st.caption(
            "下方「本次预测」为当前结果；参数或预测变化时会自动记入历史。"
            "可在历史记录中删除单条，或一键清空（不影响当前预测）。"
        )
        st.markdown("**本次预测**")
        st.dataframe(df_combo, use_container_width=True, height=360)

        hist: list = st.session_state.get(SFC_PRED_HISTORY_KEY) or []
        older = hist[1:] if len(hist) > 1 else []
        if older:
            st.markdown("**历史记录（较早的预测，可删除）**")
            c0, c1 = st.columns([3, 1])
            with c0:
                st.caption(f"共保存 {len(hist)} 条快照，下方列出除最近一次外的 {len(older)} 条。")
            with c1:
                if st.button("清空全部历史", key="sfc_hist_clear_all", type="secondary"):
                    st.session_state[SFC_PRED_HISTORY_KEY] = []
                    st.session_state[SFC_PRED_HIST_SIG_KEY] = (
                        self._prediction_snapshot_sig(user_inputs, flat_display)
                    )
                    st.rerun()

            for rec in older:
                title = f"{rec['ts']}  ·  id={rec['id']}"
                with st.expander(title, expanded=False):
                    oin = pd.DataFrame(
                        [{f"输入.{k}": v for k, v in rec["user_inputs"].items()}]
                    )
                    oout = pd.DataFrame([rec["flat_display"]])
                    st.dataframe(
                        pd.concat([oin, oout], axis=1),
                        use_container_width=True,
                        height=280,
                    )
                    if st.button(
                        "删除此条",
                        key=f"sfc_hist_del_{rec['id']}",
                        type="secondary",
                    ):
                        st.session_state[SFC_PRED_HISTORY_KEY] = [
                            x
                            for x in st.session_state[SFC_PRED_HISTORY_KEY]
                            if x["id"] != rec["id"]
                        ]
                        st.session_state[SFC_PRED_HIST_SIG_KEY] = (
                            self._prediction_snapshot_sig(user_inputs, flat_display)
                        )
                        st.rerun()
        else:
            if len(hist) == 1:
                st.caption(
                    "当前历史仅保存了 1 条快照（与上方「本次预测」一致）。"
                    "调整侧栏参数使预测变化后，较早的快照将出现在下方并可逐条删除。"
                )
                if st.button("清空全部历史", key="sfc_hist_clear_single", type="secondary"):
                    st.session_state[SFC_PRED_HISTORY_KEY] = []
                    st.session_state[SFC_PRED_HIST_SIG_KEY] = (
                        self._prediction_snapshot_sig(user_inputs, flat_display)
                    )
                    st.rerun()
            else:
                st.caption(
                    "暂无历史快照，或已全部清空。上方仍为当前预测；参数或结果变化后将自动写入历史。"
                )

    def show_mechanism_analysis(
        self,
        predictor,
        X_df: pd.DataFrame,
        pred_result: dict,
        *,
        include_offline_json: bool = True,
    ) -> None:
        """高级分析：特征重要性、PDP、ICE、完整 SHAP（研发向）。"""
        st.markdown(
            section_title_html("特征趋势与模型诊断"),
            unsafe_allow_html=True,
        )
        st.caption(
            "PDP 为训练样本上的平均边际趋势；ICE 为固定当前工况的单特征扫描。"
            "仅供研发理解，不等同于力学解析。"
        )

        mid = pred_result.get("intermediate") or {}
        pred_class = int(mid.get("risk_class_argmax", 0))

        imp_path = MODELS_DIR / "feature_importance.json"
        imp_all = load_feature_importance_json(MODELS_DIR)

        if include_offline_json:
            with st.expander("模型训练指标（单次划分·测试集）", expanded=False):
                tm = load_training_metrics_json(MODELS_DIR)
                if tm:
                    st.json(tm)
                else:
                    st.caption("未找到 `training_metrics.json`。")

            with st.expander("K 折交叉验证（全数据·论文级）", expanded=False):
                cv = load_cv_metrics_json(MODELS_DIR)
                if cv:
                    st.caption(
                        f"来源 `models/cv_metrics.json`：{cv.get('n_splits', '?')} 折，"
                        f"n={cv.get('n_samples', '?')}。"
                    )
                    st.json(cv)
                else:
                    st.caption("未找到 `models/cv_metrics.json`。")

        if not imp_all:
            st.warning(
                f"未找到 `{imp_path.name}`，无法进行特征重要性展示。"
                "请在本项目目录执行训练或删除 `models` 下 pkl 以触发演示模型重建。"
            )
            return

        with st.expander("全局特征重要性（gain）", expanded=False):
            self._render_feature_importance_tabs(imp_all)

        # 背景与当前点（标准化空间）
        Xo = X_df[FEATURE_COLUMNS].astype(np.float64, copy=False)
        x_scaled = predictor.scaler.transform(Xo)
        x_single = x_scaled[0:1]
        X_bg = load_background_scaled(predictor.scaler)
        if X_bg is None:
            X_bg = synthetic_background_around(x_single, n=220)
            st.caption(
                "未读取到 `data/training_data.example.csv`，已用当前工况加小扰动合成背景样本，PDP 为近似。"
            )

        w_imp = imp_all.get("crack_width") or {}
        top_list = top_feature_names(w_imp, n=14) if w_imp else FEATURE_COLUMNS[:12]
        default_feats = ["w_b_ratio", "fiber_content"]
        trend_feats = [f for f in default_feats if f in top_list]
        if not trend_feats:
            trend_feats = top_list[:2]
        else:
            for f in top_list:
                if f not in trend_feats and len(trend_feats) < 2:
                    trend_feats.append(f)

        st.markdown(
            section_title_html("特征趋势分析（高级）"),
            unsafe_allow_html=True,
        )
        st.caption(
            f"默认展示：{'、'.join(self._feat_label_plain(f) for f in trend_feats)}。"
            "坐标为模型内部标准化空间。"
        )
        t1, t2 = st.columns(2)
        for col, feat_name in zip((t1, t2), trend_feats[:2]):
            with col:
                self._render_pdp_ice_pair_for_feature(
                    predictor,
                    X_bg,
                    x_single,
                    pred_class,
                    feat_name,
                    compact=True,
                )

        with st.expander("更多特征 · PDP / ICE / 自定义", expanded=False):
            feat_pick = st.selectbox(
                "选择特征",
                options=top_list,
                format_func=lambda n: self._feat_label_plain(n),
                index=0,
                key="mech_feature_pick_adv",
            )
            self._render_pdp_ice_pair_for_feature(
                predictor,
                X_bg,
                x_single,
                pred_class,
                feat_pick,
                compact=False,
            )

        with st.expander("完整风险贡献图（SHAP · TreeExplainer）", expanded=False):
            self._render_shap_risk_advanced(
                predictor, X_bg, x_single, pred_class, top_k=14
            )

        st.caption(f"特征重要性数据：`{imp_path}`")

    def _render_feature_importance_tabs(self, imp_all: dict) -> None:
        task_labels = {
            "crack_width": "裂缝宽度",
            "crack_density": "裂缝密度",
            "cracking_risk": "开裂风险",
        }
        tabs = st.tabs(
            [task_labels[k] for k in ["crack_width", "crack_density", "cracking_risk"]]
        )
        for tab, key in zip(
            tabs, ["crack_width", "crack_density", "cracking_risk"], strict=True
        ):
            with tab:
                raw = imp_all.get(key) or {}
                if not raw:
                    st.caption("无此项重要性数据。")
                    continue
                names = list(raw.keys())
                vals = [float(raw[n]) for n in names]
                s = sum(abs(v) for v in vals) or 1.0
                vals_n = [abs(v) / s for v in vals]
                order = np.argsort(-np.array(vals_n))[:18]
                y_lab = [self._feat_label_plain(names[i]) for i in order][::-1]
                x_v = [vals_n[i] for i in order][::-1]
                fig = go.Figure(
                    go.Bar(
                        x=x_v,
                        y=y_lab,
                        orientation="h",
                        marker=dict(
                            color=x_v,
                            colorscale=[
                                [0, COLORS["primary_light"]],
                                [1, COLORS["primary"]],
                            ],
                            showscale=False,
                            line=dict(width=0),
                        ),
                    )
                )
                hi = max(380, 26 * len(y_lab))
                fig.update_layout(
                    title=task_labels[key],
                    xaxis_title="相对重要性",
                    margin=dict(l=168, r=28, t=56, b=44),
                )
                apply_chart_theme(fig, height=hi)
                self._plotly(fig)

    def _render_pdp_ice_pair_for_feature(
        self,
        predictor,
        X_bg: np.ndarray,
        x_single: np.ndarray,
        pred_class: int,
        feat_name: str,
        *,
        compact: bool,
    ) -> None:
        j = FEATURE_COLUMNS.index(feat_name)
        lab = self._feat_label_plain(feat_name)
        st.markdown(f"**{lab}**")
        if compact:
            c1, c2 = st.columns(2)
        else:
            st.markdown("##### 部分依赖 PDP")
            c1, c2, c3 = st.columns(3)
            d1, d2, d3 = st.columns(3)

        def _pdp_plot(col, model, title, is_clf=False):
            with col:
                if model is None:
                    st.caption(f"{title} 未加载")
                    return
                gx, gy = partial_dependence_1d(
                    model,
                    X_bg,
                    j,
                    is_classifier=is_clf,
                    class_index=pred_class if is_clf else 0,
                )
                if gx is None:
                    st.caption("PDP 跳过")
                    return
                fig_p = go.Figure(
                    go.Scatter(
                        x=gx,
                        y=gy,
                        mode="lines",
                        line=dict(color=COLORS["primary"], width=2.4),
                    )
                )
                fig_p.update_layout(
                    title=title,
                    xaxis_title=lab,
                    height=300 if compact else 340,
                )
                apply_chart_theme(fig_p, height=300 if compact else 340)
                self._plotly(fig_p)

        if compact:
            with c1:
                _pdp_plot(c1, predictor.crack_risk_model, "开裂风险 PDP", is_clf=True)
            with c2:
                if predictor.crack_risk_model is None:
                    st.caption("ICE 未加载")
                else:
                    grid, yv = local_ice_curve_proba(
                        predictor.crack_risk_model,
                        x_single,
                        X_bg,
                        j,
                        pred_class,
                    )
                    fig_i = go.Figure(
                        go.Scatter(
                            x=grid,
                            y=yv,
                            mode="lines",
                            line=dict(color=COLORS["warn"], width=2.4),
                        )
                    )
                    fig_i.update_layout(
                        title="开裂风险 ICE",
                        xaxis_title=lab,
                        height=300,
                    )
                    apply_chart_theme(fig_i, height=300)
                    self._plotly(fig_i)
            return

        _pdp_plot(c1, predictor.crack_width_model, "裂缝宽度")
        _pdp_plot(c2, predictor.crack_density_model, "裂缝密度")
        _pdp_plot(c3, predictor.crack_risk_model, "开裂风险", is_clf=True)
        st.markdown("##### 局部 ICE")
        with d1:
            if predictor.crack_width_model is None:
                st.caption("裂缝宽度未加载")
            else:
                grid, yv = local_ice_curve(
                    predictor.crack_width_model, x_single, X_bg, j
                )
                fig_i = go.Figure(
                    go.Scatter(
                        x=grid, y=yv, mode="lines",
                        line=dict(color=COLORS["success"], width=2.4),
                    )
                )
                fig_i.update_layout(title="裂缝宽度", xaxis_title=lab, height=320)
                apply_chart_theme(fig_i, height=320)
                self._plotly(fig_i)
        with d2:
            if predictor.crack_density_model is None:
                st.caption("裂缝密度未加载")
            else:
                grid, yv = local_ice_curve(
                    predictor.crack_density_model, x_single, X_bg, j
                )
                fig_i = go.Figure(
                    go.Scatter(
                        x=grid, y=yv, mode="lines",
                        line=dict(color=COLORS["success"], width=2.4),
                    )
                )
                fig_i.update_layout(title="裂缝密度", xaxis_title=lab, height=320)
                apply_chart_theme(fig_i, height=320)
                self._plotly(fig_i)
        with d3:
            if predictor.crack_risk_model is None:
                st.caption("开裂风险未加载")
            else:
                grid, yv = local_ice_curve_proba(
                    predictor.crack_risk_model, x_single, X_bg, j, pred_class
                )
                fig_i = go.Figure(
                    go.Scatter(
                        x=grid, y=yv, mode="lines",
                        line=dict(color=COLORS["warn"], width=2.4),
                    )
                )
                fig_i.update_layout(title="开裂风险", xaxis_title=lab, height=320)
                apply_chart_theme(fig_i, height=320)
                self._plotly(fig_i)

    def _render_shap_risk_advanced(
        self,
        predictor,
        X_bg: np.ndarray,
        x_single: np.ndarray,
        pred_class: int,
        *,
        top_k: int = 14,
    ) -> None:
        if predictor.crack_risk_model is None:
            st.info("开裂风险模型未加载，无法计算 SHAP。")
            return
        try:
            shp = shap_risk_bar(
                predictor.crack_risk_model,
                X_bg,
                x_single,
                pred_class,
                top_k=top_k,
            )
        except Exception:
            shp = None
        if not shp:
            st.info("SHAP（TreeExplainer）未可用，已跳过。")
            return
        names, values = shp
        labels = [f"{self._feat_label_plain(n)} ({n})" for n in names]
        colors = np.where(values >= 0, COLORS["danger"], COLORS["primary_light"])
        fig_s = go.Figure(
            go.Bar(
                x=values,
                y=labels,
                orientation="h",
                marker=dict(color=colors, line=dict(width=0)),
            )
        )
        h_shap = max(420, 24 * len(labels))
        fig_s.update_layout(
            title=f"TreeExplainer · 类别 {pred_class}",
            xaxis_title="SHAP",
            height=h_shap,
            margin=dict(l=200, r=28, t=48, b=44),
        )
        apply_chart_theme(fig_s, height=h_shap)
        self._plotly(fig_s)

    def show_sensitivity(self, result: dict) -> None:
        st.markdown(
            section_title_html("参数敏感性（规划）"),
            unsafe_allow_html=True,
        )
        st.info("参数敏感性分析功能将基于多次扰动模拟实现。")

    def _render_one_recommendation(self, rec: dict) -> None:
        title = rec.get("title", "建议")
        with st.expander(title, expanded=True):
            if rec.get("suggestion"):
                st.write(rec["suggestion"])
            if rec.get("expected_improvement"):
                st.write("预期改善效果：" + rec["expected_improvement"])
            if rec.get("cost_impact"):
                st.write("成本影响：" + rec["cost_impact"])

    def show_optimization_suggestions(self, result: dict) -> None:
        st.markdown(section_title_html("优化建议"), unsafe_allow_html=True)
        raw = result.get("recommendations")
        if raw is None:
            recs = []
        elif isinstance(raw, list):
            recs = [r for r in raw if isinstance(r, dict)]
        else:
            recs = []

        def _src(r: dict) -> str | None:
            s = r.get("source")
            if s is None:
                return None
            return str(s).strip()

        pred_recs = [r for r in recs if _src(r) == "prediction"]
        rule_recs = [r for r in recs if _src(r) != "prediction"]

        if not pred_recs and not rule_recs:
            st.warning(
                "未读取到任何优化建议条目。常见原因：① 浏览器/Streamlit 使用了旧版脚本缓存；"
                "② 预测未完成或 `predict_all` 未写入 `recommendations`。"
            )
            st.write("暂无优化建议。")
            st.caption(
                "请尝试：右上角 **⋮** → **Clear cache** → **Clear all**，保存代码后重新运行 "
                "`streamlit run app.py`，并确认工作目录为 **SteelFiberCrackPredictor** 项目根。"
            )
            return

        if pred_recs:
            st.markdown(
                section_title_html("根据本次预测结果"),
                unsafe_allow_html=True,
            )
            st.caption(
                "以下条目由当前预测得到的时间维度、状态维度与风险概率 P 等自动触发，可与下方规则类建议一并参考。"
            )
            for rec in pred_recs:
                self._render_one_recommendation(rec)

        if rule_recs:
            st.markdown(
                section_title_html("根据配合比、材料与工程类型规则"),
                unsafe_allow_html=True,
            )
            st.caption("在纤维材质、强度等级、混凝土类型等条件满足时由规则库命中。")
            for rec in rule_recs:
                self._render_one_recommendation(rec)

    def _write_doc_lab_section(self, doc, user_inputs: dict) -> None:
        """Word：三、试验估算（抗压/抗折）。"""
        doc.add_heading("三、试验估算（抗压/抗折）", level=2)
        doc.add_paragraph(
            self._word_safe_text(
                "本节与系统「试验估算」页一致：抗压/抗折给出国标导向的公式基线 compressive_formula_pred、"
                "flexural_formula_pred（与 compressive_mpa、flexural_mpa 数值一致）；"
                "不能替代 GB/T 50081 等标准试验与正式检测报告。"
            )
        )

        grade = user_inputs.get("strength_grade", "C30")
        try:
            fcu = float(STRENGTH_GRADE_TO_MPA[grade])
        except (KeyError, TypeError, ValueError):
            fcu = 30.0
        try:
            fc_pct = float(user_inputs.get("fiber_content", 1.0))
        except (TypeError, ValueError):
            fc_pct = 1.0

        def _g(key: str, default):
            return st.session_state.get(key, default)

        scope = _g("lab_estimate_scope", "抗压与抗折")
        compute_flex = scope == "抗压与抗折"

        spec = _g("lab_specimen", SPECIMEN_TYPES[0])
        load_c = _g("lab_load_c", LOADING_COMPRESSION[0])
        load_f = (
            _g("lab_load_f", LOADING_FLEXURAL[0]) if compute_flex else None
        )
        cube_edge = float(_g("lab_cube_a", 150.0))
        pb = float(_g("lab_pb", 150.0))
        ph = float(_g("lab_ph", 150.0))
        pl = float(_g("lab_pl", 300.0))
        bb = float(_g("lab_bb", 100.0))
        bh = float(_g("lab_bh", 100.0))
        bspan = float(_g("lab_span", 400.0))

        try:
            est = estimate_strengths(
                spec,
                cube_edge_mm=cube_edge,
                prism_b_mm=pb,
                prism_h_mm=ph,
                prism_l_mm=pl,
                beam_b_mm=bb,
                beam_h_mm=bh,
                beam_span_mm=bspan,
                loading_compression=load_c,
                loading_flexural=load_f,
                cube_strength_mpa=fcu,
                fiber_content_pct=fc_pct,
                compute_flexural=compute_flex,
            )
        except Exception as ex:
            doc.add_paragraph(
                self._word_safe_text(
                    f"试验结果计算失败：{type(ex).__name__}: {ex}"
                )
            )
            return

        tbl = doc.add_table(rows=1, cols=2)
        self._table_grid_safe(tbl)
        tbl.rows[0].cells[0].text = self._word_safe_text("项目")
        tbl.rows[0].cells[1].text = self._word_safe_text("取值")
        lab_rows: list[tuple[str, str]] = [
            ("估算范围", str(scope)),
            ("强度等级（侧栏）", str(grade)),
            ("纤维体积掺量（%）", f"{fc_pct:.2f}"),
            ("试件类型", str(spec)),
            ("抗压试验加载方式", str(load_c)),
            (
                "抗折试验加载方式",
                str(load_f) if load_f is not None else "（本次未估算抗折）",
            ),
        ]
        if spec == "立方体（边长可变）":
            lab_rows.append(("立方体边长 a（mm）", f"{cube_edge:.0f}"))
        elif spec == "棱柱体（轴心抗压）":
            lab_rows.append(
                ("棱柱体 b×h×L（mm）", f"{pb:.0f}×{ph:.0f}×{pl:.0f}")
            )
        else:
            lab_rows.append(
                ("梁 b×h×跨度（mm）", f"{bb:.0f}×{bh:.0f}×{bspan:.0f}")
            )
        for label, val in lab_rows:
            row = tbl.add_row().cells
            row[0].text = self._word_safe_text(label)
            row[1].text = self._word_safe_text(val)

        doc.add_heading("（一）估算结果", level=3)
        t2 = doc.add_table(rows=1, cols=2)
        self._table_grid_safe(t2)
        t2.rows[0].cells[0].text = self._word_safe_text("指标")
        t2.rows[0].cells[1].text = self._word_safe_text("数值")
        flex_disp = (
            f"{est.flexural_formula_pred:.2f}"
            if compute_flex and math.isfinite(est.flexural_formula_pred)
            else "—"
        )
        for label, val in [
            ("抗压强度·公式基线 compressive_formula_pred（MPa）", f"{est.compressive_formula_pred:.2f}"),
            ("抗折强度·公式基线 flexural_formula_pred（MPa）", flex_disp),
        ]:
            row = t2.add_row().cells
            row[0].text = self._word_safe_text(label)
            row[1].text = self._word_safe_text(val)

        doc.add_heading("（二）中间量与说明", level=3)
        for k, v in est.detail.items():
            doc.add_paragraph(
                self._word_safe_text(f"{k}：{v}")
            )
        for note in est.notes:
            doc.add_paragraph(self._word_safe_text(note))

    def _write_doc_mechanism_section(
        self,
        doc,
        result: dict,
        predictor,
        X_df: pd.DataFrame | None,
    ) -> None:
        """Word：四、机理分析（特征重要性、测试指标、SHAP 摘要）。"""
        doc.add_heading("四、机理分析（数据驱动）", level=2)
        doc.add_paragraph(
            self._word_safe_text(
                "本节汇总当前 XGBoost 模型在训练阶段给出的特征相对重要性（基于 gain、表中为归一化值），"
                "并插入与网页「机理分析」一致的条形图、部分依赖（PDP）、局部响应（ICE）及开裂风险 SHAP 附图；"
                "图下附简要文字解读。完整交互式曲线仍可在系统「机理分析」页查看。"
            )
        )

        tm = load_training_metrics_json(MODELS_DIR)
        doc.add_heading("（一）模型测试集表现摘要", level=3)
        if tm:
            w = tm.get("crack_width") or {}
            d = tm.get("crack_density") or {}
            r = tm.get("cracking_risk") or {}
            doc.add_paragraph(
                self._word_safe_text(
                    f"裂缝宽度：MAE={self._fmt_value(w.get('test_mae'))} mm，"
                    f"RMSE={self._fmt_value(w.get('test_rmse'))}，R2={self._fmt_value(w.get('test_r2'))}。"
                )
            )
            doc.add_paragraph(
                self._word_safe_text(
                    f"裂缝密度：MAE={self._fmt_value(d.get('test_mae'))}，"
                    f"RMSE={self._fmt_value(d.get('test_rmse'))}，R2={self._fmt_value(d.get('test_r2'))}。"
                )
            )
            doc.add_paragraph(
                self._word_safe_text(
                    f"开裂风险（多分类）：准确率={self._fmt_value(r.get('test_accuracy'))}，"
                    f"宏平均 F1={self._fmt_value(r.get('test_macro_f1'))}。"
                )
            )
            doc.add_paragraph(
                self._word_safe_text(
                    f"训练/测试划分样本量：训练 {tm.get('n_train', '—')}，测试 {tm.get('n_test', '—')}。"
                )
            )
        else:
            doc.add_paragraph(
                self._word_safe_text(
                    "未找到 training_metrics.json，略过测试集指标（可运行训练脚本生成）。"
                )
            )

        imp_all = load_feature_importance_json(MODELS_DIR)
        doc.add_heading("（二）特征重要性 Top 10（归一化）", level=3)
        task_cn_map = {
            "crack_width": "裂缝宽度回归",
            "crack_density": "裂缝密度回归",
            "cracking_risk": "开裂风险分类",
        }
        if imp_all:
            for key, title in [
                ("crack_width", "1）裂缝宽度回归"),
                ("crack_density", "2）裂缝密度回归"),
                ("cracking_risk", "3）开裂风险分类"),
            ]:
                raw = imp_all.get(key) or {}
                rows = top_feature_importance_rows(raw, 10)
                if not rows:
                    continue
                doc.add_paragraph(self._word_safe_text(title))
                tbl = doc.add_table(rows=1, cols=3)
                self._table_grid_safe(tbl)
                c0, c1, c2 = tbl.rows[0].cells
                c0.text = self._word_safe_text("特征（中文）")
                c1.text = self._word_safe_text("特征（英文键）")
                c2.text = self._word_safe_text("相对重要性")
                for cn, en, val in rows:
                    row = tbl.add_row().cells
                    row[0].text = self._word_safe_text(cn)
                    row[1].text = self._word_safe_text(en)
                    try:
                        row[2].text = self._word_safe_text(f"{float(val):.4f}")
                    except (TypeError, ValueError):
                        row[2].text = "—"
                try:
                    from docx.shared import Inches

                    from src.report_figures import (
                        importance_bar_png_bytes,
                        text_analysis_importance,
                    )

                    labels = [r[0] for r in rows]
                    vals = [float(r[2]) for r in rows]
                    fig_title = f"{title} — 特征重要性（条形图）"
                    buf = importance_bar_png_bytes(labels, vals, fig_title)
                    doc.add_picture(buf, width=Inches(6.2))
                    doc.add_paragraph(
                        self._word_safe_text(
                            text_analysis_importance(
                                task_cn_map.get(key, title), labels, vals
                            )
                        )
                    )
                except Exception as ex:
                    doc.add_paragraph(
                        self._word_safe_text(
                            f"特征重要性附图生成跳过：{type(ex).__name__}: {ex}"
                        )
                    )
        else:
            doc.add_paragraph(
                self._word_safe_text(
                    "未找到 feature_importance.json，无法进行特征重要性摘要。"
                )
            )

        doc.add_heading("（三）部分依赖（PDP）与局部响应（ICE）", level=3)
        if predictor is None or X_df is None:
            doc.add_paragraph(
                self._word_safe_text(
                    "未传入模型与特征矩阵，略过 PDP/ICE 附图（请从预测流程生成报告以包含此项）。"
                )
            )
        else:
            try:
                from docx.shared import Inches

                from src.report_figures import (
                    ice_line_png_bytes,
                    pdp_line_png_bytes,
                    text_analysis_ice,
                    text_analysis_pdp,
                )

                mid = result.get("intermediate") or {}
                pred_class = int(mid.get("risk_class_argmax", 0))
                Xo = X_df[FEATURE_COLUMNS].astype(np.float64, copy=False)
                x_s = predictor.scaler.transform(Xo)
                X_bg = load_background_scaled(predictor.scaler)
                if X_bg is None:
                    X_bg = synthetic_background_around(x_s[0:1], n=220)
                w_imp = (imp_all or {}).get("crack_width") or {}
                top_names = top_feature_names(w_imp, n=1) if w_imp else []
                feat_pick = top_names[0] if top_names else FEATURE_COLUMNS[0]
                j = FEATURE_COLUMNS.index(feat_pick)
                feat_lab = feature_label(feat_pick)
                doc.add_paragraph(
                    self._word_safe_text(
                        f"以下 PDP/ICE 均针对「裂缝宽度模型」重要性最高的特征：{feat_lab}（{feat_pick}），"
                        "与网页机理分析默认一致；横轴为标准化特征空间。"
                    )
                )
                x_ax = f"{feat_lab}（标准化）"

                if predictor.crack_width_model is not None:
                    gx, gy = partial_dependence_1d(
                        predictor.crack_width_model, X_bg, j, is_classifier=False
                    )
                    if gx is not None and gy is not None:
                        buf = pdp_line_png_bytes(
                            gx,
                            gy,
                            title="裂缝宽度 (mm) — PDP",
                            xlabel=x_ax,
                            ylabel="平均预测宽度",
                        )
                        doc.add_picture(buf, width=Inches(6.0))
                        doc.add_paragraph(
                            self._word_safe_text(
                                text_analysis_pdp(feat_lab, "裂缝宽度", gy, gx)
                            )
                        )
                    else:
                        doc.add_paragraph(self._word_safe_text("裂缝宽度 PDP 未计算成功。"))
                else:
                    doc.add_paragraph(self._word_safe_text("裂缝宽度模型未加载，略过宽度 PDP。"))

                if predictor.crack_density_model is not None:
                    gx, gy = partial_dependence_1d(
                        predictor.crack_density_model, X_bg, j, is_classifier=False
                    )
                    if gx is not None and gy is not None:
                        buf = pdp_line_png_bytes(
                            gx,
                            gy,
                            title="裂缝密度 (条/m²) — PDP",
                            xlabel=x_ax,
                            ylabel="平均预测密度",
                        )
                        doc.add_picture(buf, width=Inches(6.0))
                        doc.add_paragraph(
                            self._word_safe_text(
                                text_analysis_pdp(feat_lab, "裂缝密度", gy, gx)
                            )
                        )
                    else:
                        doc.add_paragraph(self._word_safe_text("裂缝密度 PDP 未计算成功。"))
                else:
                    doc.add_paragraph(self._word_safe_text("裂缝密度模型未加载。"))

                if predictor.crack_risk_model is not None:
                    gx, gy = partial_dependence_1d(
                        predictor.crack_risk_model,
                        X_bg,
                        j,
                        is_classifier=True,
                        class_index=pred_class,
                    )
                    if gx is not None and gy is not None:
                        buf = pdp_line_png_bytes(
                            gx,
                            gy,
                            title=f"开裂风险：类别 {pred_class} 概率 — PDP",
                            xlabel=x_ax,
                            ylabel="平均预测概率",
                        )
                        doc.add_picture(buf, width=Inches(6.0))
                        doc.add_paragraph(
                            self._word_safe_text(
                                text_analysis_pdp(
                                    feat_lab, f"类别 {pred_class} 概率", gy, gx
                                )
                            )
                        )
                    else:
                        doc.add_paragraph(
                            self._word_safe_text("开裂风险 PDP 未计算成功。")
                        )
                else:
                    doc.add_paragraph(self._word_safe_text("开裂风险模型未加载。"))

                if predictor.crack_width_model is not None:
                    grid, yv = local_ice_curve(
                        predictor.crack_width_model, x_s, X_bg, j
                    )
                    buf = ice_line_png_bytes(
                        grid,
                        yv,
                        title="裂缝宽度 — 局部响应（ICE 风格）",
                        xlabel=x_ax,
                        ylabel="预测宽度 (mm)",
                    )
                    doc.add_picture(buf, width=Inches(6.0))
                    doc.add_paragraph(
                        self._word_safe_text(
                            text_analysis_ice(
                                feat_lab, "裂缝宽度预测", yv, grid
                            )
                        )
                    )
                else:
                    doc.add_paragraph(self._word_safe_text("略过局部 ICE（宽度模型未加载）。"))
            except Exception as ex:
                doc.add_paragraph(
                    self._word_safe_text(
                        f"PDP/ICE 附图写入异常：{type(ex).__name__}: {ex}"
                    )
                )

        doc.add_heading("（四）当前样本：开裂风险 SHAP", level=3)
        if predictor is None or X_df is None:
            doc.add_paragraph(
                self._word_safe_text(
                    "未传入模型与特征矩阵，略过 SHAP（请从预测流程生成报告以包含此项）。"
                )
            )
            return
        try:
            from docx.shared import Inches

            from src.report_figures import shap_bar_png_bytes, text_analysis_shap

            mid = result.get("intermediate") or {}
            pred_class = int(mid.get("risk_class_argmax", 0))
            Xo = X_df[FEATURE_COLUMNS].astype(np.float64, copy=False)
            x_s = predictor.scaler.transform(Xo)
            X_bg = load_background_scaled(predictor.scaler)
            if X_bg is None:
                X_bg = synthetic_background_around(x_s[0:1], n=200)
            if predictor.crack_risk_model is None:
                doc.add_paragraph(self._word_safe_text("开裂风险模型未加载。"))
                return
            shp = shap_risk_bar(
                predictor.crack_risk_model,
                X_bg,
                x_s,
                pred_class,
                top_k=12,
            )
            if shp:
                names, values = shp
                labs_cn = [feature_label(en) for en in names]
                buf = shap_bar_png_bytes(
                    labs_cn,
                    values,
                    f"开裂风险 SHAP（类别 {pred_class}，Top 12）",
                )
                doc.add_picture(buf, width=Inches(6.2))
                doc.add_paragraph(
                    self._word_safe_text(text_analysis_shap(pred_class, labs_cn, values))
                )
                doc.add_paragraph(
                    self._word_safe_text(
                        f"下表为与上图对应的数值（TreeExplainer）。正值表示该特征推高类别 {pred_class} 的概率，负值表示拉低。"
                    )
                )
                tbl = doc.add_table(rows=1, cols=3)
                self._table_grid_safe(tbl)
                h = tbl.rows[0].cells
                h[0].text = self._word_safe_text("序号")
                h[1].text = self._word_safe_text("特征")
                h[2].text = self._word_safe_text("SHAP")
                for i, (en, val) in enumerate(zip(names, values), start=1):
                    row = tbl.add_row().cells
                    row[0].text = self._word_safe_text(str(i))
                    row[1].text = self._word_safe_text(
                        f"{feature_label(en)}（{en}）"
                    )
                    try:
                        row[2].text = self._word_safe_text(f"{float(val):.6f}")
                    except (TypeError, ValueError):
                        row[2].text = "—"
            else:
                doc.add_paragraph(
                    self._word_safe_text(
                        "SHAP 未成功计算（可能与 shap / XGBoost 版本有关），请以网页「机理分析」为准。"
                    )
                )
        except Exception as ex:
            doc.add_paragraph(
                self._word_safe_text(f"SHAP 摘要写入异常：{type(ex).__name__}: {ex}")
            )

    def _render_lab_strength_offline_metrics_body(self) -> None:
        """只读 lab_strength_residual_report.json 正文（无外层 expander）。"""
        report_path = OUTPUTS_DIR / "lab_strength" / "lab_strength_residual_report.json"
        zh_disclaimer = (
            "以下指标来自历史训练数据的离线评估，用于说明历史拟合能力，"
            "不代表当前输入条件下的真实试验误差。\n\n"
            "当输入参数超出训练分布，或试件类型、龄期、加载方式与训练数据不一致时，估算结果仅供参考。"
        )
        st.markdown(zh_disclaimer)
        st.divider()
        if not report_path.is_file():
            st.info(
                "未找到离线训练报告文件：`outputs/lab_strength/lab_strength_residual_report.json`。"
                "请在项目根目录运行 `train_lab_strength_residual.py` 生成报告后，再查看下方汇总指标。"
            )
            return
        try:
            with open(report_path, encoding="utf-8") as f:
                rep = json.load(f)
        except (OSError, json.JSONDecodeError) as e:
            st.info(
                f"无法读取训练报告 JSON（{type(e).__name__}）。"
                "请确认文件为有效 UTF-8 JSON。"
            )
            return

        dm = rep.get("default_method_by_task")
        comp_dm = dm.get("compressive") if isinstance(dm, dict) else None
        flex_dm = dm.get("flexural") if isinstance(dm, dict) else None
        comp_task = rep.get("compressive")
        flex_task = rep.get("flexural")

        st.markdown("##### 公式说明")
        st.markdown(
            """
本页试验估算结果以公式基线为核心，用于给出规范路径下的参数化参考值，并非无约束黑箱拟合。

离线训练报告中的指标，反映的是历史训练数据上的交叉验证表现，不等同于当前单次输入的实时预测误差。

因此，本页展示值应优先理解为当前参数条件下的公式估算结果。
            """.strip()
        )

        def _default_path_cn(label: str, block: object) -> tuple[str | None, str | None]:
            """从 JSON 子块生成一行可读说明；缺失时返回 (None, caption)。"""
            if not isinstance(block, dict):
                return None, f"{label}：离线报告中缺少 `default_method_by_task` 下对应子对象。"
            strat = block.get("strategy")
            learner = block.get("residual_learner")
            if strat is None and (
                learner is None or (isinstance(learner, str) and not learner.strip())
            ):
                return None, f"{label}：缺少 `strategy` 与 `residual_learner` 字段，无法生成摘要。"
            ln = ""
            if isinstance(learner, str) and learner.strip():
                lk = learner.strip().lower()
                ln = {"ridge": "Ridge", "linear": "线性", "hgb": "梯度提升"}.get(
                    lk, learner.strip()
                )
            if ln:
                tail = f"（报告字段 `strategy={strat!r}`，`residual_learner={learner!r}`）"
                return f"{label}当前默认：**公式基线 + {ln} 残差修正**{tail}", None
            if isinstance(strat, str) and "formula_only" in strat.lower():
                return (
                    f"{label}当前默认：**仅公式基线**（报告字段 `strategy={strat!r}`，`residual_learner` 为空）",
                    None,
                )
            return (
                f"{label}当前默认：请结合报告字段理解（`strategy={strat!r}`，`residual_learner={learner!r}`）",
                None,
            )

        if not isinstance(dm, dict):
            st.caption(
                "离线报告中缺少默认策略字段（`default_method_by_task`），"
                "无法从 JSON 自动生成「抗压/抗折当前默认路径」摘要。"
            )
        else:
            c_line, c_warn = _default_path_cn("抗压", comp_dm)
            f_line, f_warn = _default_path_cn("抗折", flex_dm)
            if c_line:
                st.markdown(f"- {c_line}")
            elif c_warn:
                st.caption(c_warn)
            if f_line:
                st.markdown(f"- {f_line}")
            elif f_warn:
                st.caption(f_warn)

        st.divider()

        st.markdown("##### 离线报告汇总（训练数据）")
        n_final = rep.get("n_rows_final_used")
        if n_final is None:
            st.caption("报告中缺少 `n_rows_final_used`。")
        else:
            st.markdown(f"- **n_rows_final_used**：{n_final}")

        if isinstance(comp_dm, dict):
            strat = comp_dm.get("strategy")
            desc = comp_dm.get("description")
            learner = comp_dm.get("residual_learner")
            line = f"- **抗压默认方法**：`{strat}`"
            if desc:
                line += f" — {desc}"
            st.markdown(line)
            mkey_c = "formula_only"
            if isinstance(learner, str) and learner.strip():
                lk = learner.strip().lower()
                if lk in ("ridge", "linear", "hgb"):
                    mkey_c = lk
            elif isinstance(strat, str) and "ridge" in strat.lower():
                mkey_c = "ridge"
            oof_c = None
            if isinstance(comp_task, dict):
                og = comp_task.get("oof_global_metrics")
                if isinstance(og, dict):
                    oof_c = og.get(mkey_c)
            n_c = None
            if isinstance(comp_task, dict):
                cv = comp_task.get("cv_fold_metrics")
                if isinstance(cv, dict):
                    n_c = cv.get("n_samples")
            if (
                isinstance(oof_c, dict)
                and all(
                    k in oof_c and isinstance(oof_c.get(k), (int, float))
                    for k in ("mae", "rmse", "r2")
                )
                and isinstance(n_c, (int, float))
            ):
                st.markdown(
                    f"- **抗压 OOF（{mkey_c}）**：MAE={float(oof_c['mae']):.4f}，"
                    f"RMSE={float(oof_c['rmse']):.4f}，R²={float(oof_c['r2']):.4f}，n={int(n_c)}"
                )
            else:
                st.caption(
                    "报告中缺少抗压任务 `oof_global_metrics` 或与默认方法对应的子项，"
                    "或缺少 `cv_fold_metrics.n_samples`，不展示抗压 OOF 数值。"
                )
        else:
            st.caption("报告中缺少 `default_method_by_task.compressive`。")

        if isinstance(flex_dm, dict):
            strat_f = flex_dm.get("strategy")
            desc_f = flex_dm.get("description")
            line_f = f"- **抗折默认方法**：`{strat_f}`"
            if desc_f:
                line_f += f" — {desc_f}"
            st.markdown(line_f)
            mkey_f = "formula_only"
            oof_f = None
            if isinstance(flex_task, dict):
                ogf = flex_task.get("oof_global_metrics")
                if isinstance(ogf, dict):
                    oof_f = ogf.get(mkey_f)
            n_f = None
            if isinstance(flex_task, dict):
                cvf = flex_task.get("cv_fold_metrics")
                if isinstance(cvf, dict):
                    n_f = cvf.get("n_samples")
            if (
                isinstance(oof_f, dict)
                and all(
                    k in oof_f and isinstance(oof_f.get(k), (int, float))
                    for k in ("mae", "rmse", "r2")
                )
                and isinstance(n_f, (int, float))
            ):
                st.markdown(
                    f"- **抗折 OOF（{mkey_f}）**：MAE={float(oof_f['mae']):.4f}，"
                    f"RMSE={float(oof_f['rmse']):.4f}，R²={float(oof_f['r2']):.4f}，n={int(n_f)}"
                )
            else:
                st.caption(
                    "报告中缺少抗折任务 `oof_global_metrics.formula_only`，"
                    "或缺少 `cv_fold_metrics.n_samples`，不展示抗折 OOF 数值。"
                )
        else:
            st.caption("报告中缺少 `default_method_by_task.flexural`。")

    def show_lab_experiment(self, pred_result: dict, user_inputs: dict) -> None:
        """兼容入口：力学展示已并入 Tab「力学性能」。"""
        _ = pred_result
        self.show_tab_mechanical_support(user_inputs)

    def show_report_export(
        self,
        result: dict,
        user_inputs: dict,
        predictor=None,
        X_df: pd.DataFrame | None = None,
    ) -> None:
        st.markdown(section_title_html("报告生成"), unsafe_allow_html=True)
        st.caption(
            "Word 中含：预测三维度、**试验估算**（抗压/抗折）、**机理分析**（测试指标、特征重要性、PDP/ICE、SHAP）及优化建议。"
            "试验参数与网页「试验估算」当前选项一致。"
        )
        out_fixed = PROJECT_ROOT / "111.docx"
        st.write(
            f"点击按钮后生成 Word 报告（固定保存到项目目录）：`{out_fixed}`"
        )
        if st.button("生成预测报告（111.docx）", key="gen_predict_report_111"):
            try:
                from docx import Document

                out_path = out_fixed

                preds = result.get("predictions", {}) or {}
                recs = result.get("recommendations", []) or []
                std = result.get("standards") or {}

                doc = Document()
                doc.add_heading("纤维混凝土抗裂性能预测报告", level=1)
                doc.add_paragraph(
                    self._word_safe_text("本报告由系统根据当前输入参数自动生成。")
                )
                if std.get("disclaimer"):
                    doc.add_paragraph(self._word_safe_text(std["disclaimer"]))

                doc.add_heading("参照标准（信息性）", level=2)
                for line in std.get("referenced_specs", []):
                    doc.add_paragraph(self._word_safe_text(f"• {line}"))
                if std.get("risk_tier_note"):
                    doc.add_paragraph(self._word_safe_text(std["risk_tier_note"]))

                doc.add_heading("一、输入参数", level=2)
                table_in = doc.add_table(rows=1, cols=2)
                self._table_grid_safe(table_in)
                table_in.rows[0].cells[0].text = self._word_safe_text("参数")
                table_in.rows[0].cells[1].text = self._word_safe_text("取值")
                for k, v in user_inputs.items():
                    row = table_in.add_row().cells
                    row[0].text = self._word_safe_text(k)
                    row[1].text = self._word_safe_text(self._fmt_value(v))

                doc.add_heading("二、预测结果（三维度）", level=2)

                doc.add_heading("（一）时间维度", level=3)
                t1 = doc.add_table(rows=1, cols=2)
                self._table_grid_safe(t1)
                t1.rows[0].cells[0].text = self._word_safe_text("指标")
                t1.rows[0].cells[1].text = self._word_safe_text("数值")
                td = preds.get("time_dimension") or {}
                for label, key, unit in [
                    ("开裂时间（浇筑→首条裂缝）", "cracking_time_hours", "h"),
                    ("临界龄期", "critical_age_days", "d"),
                    ("安全窗口（养护撤除建议上限）", "safety_window_hours", "h"),
                ]:
                    row = t1.add_row().cells
                    row[0].text = self._word_safe_text(f"{label} ({unit})")
                    row[1].text = self._word_safe_text(self._fmt_value(td.get(key, "")))
                if td.get("note"):
                    doc.add_paragraph(self._word_safe_text(td["note"]))

                doc.add_heading("（二）状态维度", level=3)
                t2 = doc.add_table(rows=1, cols=2)
                self._table_grid_safe(t2)
                t2.rows[0].cells[0].text = self._word_safe_text("指标")
                t2.rows[0].cells[1].text = self._word_safe_text("数值")
                sd = preds.get("state_dimension") or {}
                for label, key in [
                    ("开裂风险概率 P（无量纲）", "risk_probability"),
                    ("最大裂缝宽度 w（mm）", "crack_width_mm"),
                    ("裂缝密度（条/m²）", "crack_density_per_m2"),
                    ("裂缝密度来源说明", "crack_density_source_cn"),
                    ("应力-强度比 σ/fₜ（启发式，无量纲）", "stress_strength_ratio"),
                ]:
                    row = t2.add_row().cells
                    row[0].text = self._word_safe_text(label)
                    row[1].text = self._word_safe_text(self._fmt_value(sd.get(key, "")))
                if sd.get("stress_strength_note_cn"):
                    doc.add_paragraph(
                        self._word_safe_text(sd["stress_strength_note_cn"])
                    )

                doc.add_heading("裂缝宽度与 GB 50010 信息性比对", level=3)
                wref = sd.get("crack_width_gb50010") or std.get(
                    "crack_width_reference_gb50010"
                )
                if isinstance(wref, dict):
                    doc.add_paragraph(
                        self._word_safe_text(
                            f"参照限值（信息性）：{wref.get('w_lim_ref_mm', '')} mm；"
                            f"{wref.get('environment_scope_cn', '')}"
                        )
                    )
                    doc.add_paragraph(
                        self._word_safe_text(wref.get("comparison_cn", ""))
                    )
                else:
                    doc.add_paragraph(self._word_safe_text("—"))

                doc.add_heading("（三）风险维度", level=3)
                rd = preds.get("risk_dimension") or {}
                doc.add_paragraph(
                    self._word_safe_text(
                        f"预警等级：{rd.get('alert_level', '')}；"
                        f"{rd.get('bands', '')}"
                    )
                )
                if rd.get("note"):
                    doc.add_paragraph(self._word_safe_text(rd["note"]))

                doc.add_heading("（四）兼容字段", level=3)
                table_out = doc.add_table(rows=1, cols=2)
                self._table_grid_safe(table_out)
                table_out.rows[0].cells[0].text = self._word_safe_text("指标")
                table_out.rows[0].cells[1].text = self._word_safe_text("预测值")
                for k, v in preds.items():
                    if k in ("time_dimension", "state_dimension", "risk_dimension"):
                        continue
                    row = table_out.add_row().cells
                    row[0].text = self._word_safe_text(k)
                    row[1].text = self._word_safe_text(self._fmt_value(v))

                try:
                    self._write_doc_lab_section(doc, user_inputs)
                except Exception as lab_err:
                    doc.add_heading("三、试验估算（抗压/抗折）", level=2)
                    doc.add_paragraph(
                        self._word_safe_text(
                            f"试验估算章节写入失败（{type(lab_err).__name__}: {lab_err}）。"
                            "请先在「试验估算」页设置试件与加载方式后重试。"
                        )
                    )

                try:
                    self._write_doc_mechanism_section(doc, result, predictor, X_df)
                except Exception as mech_err:
                    doc.add_heading("四、机理分析（数据驱动）", level=2)
                    doc.add_paragraph(
                        self._word_safe_text(
                            f"机理分析章节生成失败（{type(mech_err).__name__}: {mech_err}）。"
                            "其余章节已写入；图表请在网页「机理分析」查看。"
                        )
                    )

                doc.add_heading("五、优化建议", level=2)
                pred_doc = [r for r in recs if r.get("source") == "prediction"]
                rule_doc = [r for r in recs if r.get("source") != "prediction"]

                def _write_recs_block(items: list, start_idx: int) -> int:
                    n = start_idx
                    for rec in items:
                        title = rec.get("title", f"建议{n}")
                        doc.add_paragraph(
                            self._word_safe_text(f"{n}. {title}")
                        )
                        sug = rec.get("suggestion")
                        if sug:
                            doc.add_paragraph(
                                self._word_safe_text(f"建议内容：{sug}")
                            )
                        eff = rec.get("expected_improvement")
                        if eff:
                            doc.add_paragraph(
                                self._word_safe_text(f"预期改善：{eff}")
                            )
                        cost = rec.get("cost_impact")
                        if cost:
                            doc.add_paragraph(
                                self._word_safe_text(f"成本影响：{cost}")
                            )
                        n += 1
                    return n

                idx = 1
                if pred_doc:
                    doc.add_heading("（一）根据本次预测结果", level=3)
                    idx = _write_recs_block(pred_doc, idx)
                if rule_doc:
                    doc.add_heading(
                        "（二）根据配合比、材料与工程类型规则"
                        if pred_doc
                        else "（一）根据配合比、材料与工程类型规则",
                        level=3,
                    )
                    idx = _write_recs_block(rule_doc, idx)
                if not pred_doc and not rule_doc:
                    doc.add_paragraph(self._word_safe_text("暂无优化建议。"))

                saved_path = out_path
                try:
                    doc.save(out_path)
                except PermissionError:
                    alt = PROJECT_ROOT / (
                        f"111_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    )
                    try:
                        doc.save(alt)
                        saved_path = alt
                        st.warning(
                            self._word_safe_text(
                                "默认 111.docx 无法写入（多为 Word 正在打开该文件）。"
                                f"已另存为：{alt.name}"
                            )
                        )
                    except OSError as e2:
                        st.error(
                            self._word_safe_text(
                                f"保存失败：{e2}。请关闭 Word 中已打开的 111.docx 后重试，"
                                f"或检查目录是否可写：{PROJECT_ROOT}"
                            )
                        )
                        raise
                st.success(f"已生成：{saved_path.resolve()}")
            except Exception as e:
                import traceback

                st.error(f"生成失败：{type(e).__name__}: {e}")
                with st.expander("详细错误（可复制给开发者）", expanded=False):
                    st.code(traceback.format_exc(), language="text")

